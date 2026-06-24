# -*- coding: utf-8 -*-
"""
Genera el informe PEPI en formato Word (.docx) para el proyecto del
acueducto de Tado (Choco). Incrusta los arboles como imagenes y construye
la tabla de flujos leyendo 'resultados_modelo.csv' (consistencia total).

Requiere: python-docx
Ejecutar:  python3 generar_word.py  -> crea INFORME_PEPI.docx
"""
import os
import csv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AZUL = RGBColor(0x1F, 0x4E, 0x79)
GRIS = RGBColor(0x59, 0x59, 0x59)

doc = Document()
base = doc.styles["Normal"]
base.font.name = "Calibri"
base.font.size = Pt(11)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, white=False, size=10, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def add_table(headers, rows, col_align=None, header_fill="1F4E79", fsize=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for j, h in enumerate(headers):
        shade_cell(hdr[j], header_fill)
        set_cell_text(hdr[j], h, bold=True, white=True, size=fsize, align="center")
    for r in rows:
        cells = t.add_row().cells
        for j, val in enumerate(r):
            al = col_align[j] if col_align else "left"
            set_cell_text(cells[j], val, size=fsize, align=al)
    doc.add_paragraph()
    return t


def h1(text):
    p = doc.add_heading(level=1)
    run = p.add_run(text)
    run.font.color.rgb = AZUL
    run.font.size = Pt(15)


def h2(text):
    p = doc.add_heading(level=2)
    run = p.add_run(text)
    run.font.color.rgb = AZUL
    run.font.size = Pt(12)


def para(text, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_image_centered(path, width_inches=6.3):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))
    else:
        para("[No se encontro la imagen: " + path + ". Ejecuta generar_arboles.py]", italic=True, size=9)


def fmt(x):
    """Formatea numero con separador de miles y sin decimales; '—' si vacio."""
    s = str(x).strip()
    if s == "" or s.lower() == "none":
        return "—"
    try:
        v = float(s)
        return f"{v:,.0f}".replace(",", ".")
    except ValueError:
        return s


def leer_flujos(path="resultados_modelo.csv"):
    """Lee la seccion FLUJOS DE CAJA del CSV y devuelve filas para la tabla Word."""
    if not os.path.exists(path):
        return []
    with open(path, encoding="utf-8") as f:
        rows = list(csv.reader(f))
    # localizar encabezado de la seccion de flujos
    start = None
    for i, r in enumerate(rows):
        if r and r[0].strip() == "Anio":
            start = i + 1
            break
    if start is None:
        return []
    out = []
    for r in rows[start:]:
        if not r or r[0].strip() == "" or not r[0].strip().lstrip("-").isdigit():
            break
        # cols CSV: 0 Anio,1 Ing,2 OPEX,3 EBITDA,4 Dep,5 Int,6 Amort,7 Saldo,
        #           8 FC_Proy,9 FC_Inv,10 FC_Banco,11 Benef,12 FC_Social
        out.append([
            r[0].strip(),       # Anio
            fmt(r[1]),          # Ingresos
            fmt(r[2]),          # OPEX
            fmt(r[3]),          # EBITDA
            fmt(r[8]),          # FC Proyecto
            fmt(r[9]),          # FC Inversionista
            fmt(r[10]),         # FC Banco
            fmt(r[11]),         # Beneficio social
            fmt(r[12]),         # FC Social
        ])
    return out


# ===================== PORTADA =====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("INFORME DE PREPARACIÓN Y EVALUACIÓN\nDE PROYECTOS DE INGENIERÍA (PEPI)")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = AZUL

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Optimización y ampliación del sistema de acueducto\ndel municipio de Tadó (Chocó)")
rs.font.size = Pt(13)
rs.font.color.rgb = GRIS
doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn = note.add_run("Caso real. Cifras monetarias en millones de pesos colombianos (COP MM) salvo indicación contraria.")
rn.italic = True
rn.font.size = Pt(9)
doc.add_paragraph()

# ===================== 1. RESUMEN EJECUTIVO =====================
h1("1. Resumen Ejecutivo")
para("Tadó es un municipio del occidente del Chocó, situado sobre el río San Juan, con 17.000 habitantes según "
     "el Censo Nacional de Población y Vivienda 2018 (DANE) y uno de los regímenes de lluvia más altos del planeta "
     "(~7.900 mm/año). Pese a esa abundancia de agua, la población ha carecido históricamente de agua potable "
     "continua y de calidad: baja cobertura, discontinuidad en el servicio y un tratamiento insuficiente frente a "
     "la alta turbiedad del agua cruda durante las crecientes.")
para("El proyecto consiste en optimizar y ampliar el sistema de acueducto de la cabecera urbana —incluida una PTAP "
     "dimensionada para 45 L/s según la Resolución 0330 de 2017 (RAS)— para alcanzar continuidad 24 horas y "
     "cobertura cercana al 98 %, en línea con el proyecto real ejecutado por el Ministerio de Vivienda (MinVivienda) "
     "con cooperación española, cuya inversión fue de COP 19.971 millones.")
para("Como ocurre con casi todos los acueductos municipales pequeños en Colombia, el proyecto NO es rentable solo "
     "con tarifas (VPN financiero negativo). Su justificación es socioeconómica: con la tasa social del 9 % (DNP) el "
     "VPN es ampliamente positivo y la relación beneficio/costo es 1,56. Por eso se cofinancia con recursos públicos "
     "(SGR, Plan Departamental de Aguas, SGP y cooperación), exactamente como sucedió en Tadó.")
add_table(
    ["Indicador", "Valor"],
    [["Caudal de diseño (QMD, Res. 0330)", "41,5 L/s → PTAP de 45 L/s"],
     ["CAPEX total (anclado a inversión real)", "COP 19.971 MM"],
     ["Estructura de financiación", "80 % público / 10 % deuda / 10 % equity"],
     ["WACC", "9,31 %"],
     ["VPN del proyecto (solo tarifas) @WACC", "−15.487 MM (negativo)"],
     ["TIR del inversionista (con aporte público)", "10,46 %"],
     ["VPN socioeconómico @ tasa social 9 % (DNP)", "+19.626 MM"],
     ["Relación Beneficio/Costo (social)", "1,56"]],
    col_align=["left", "right"], fsize=10)

# ===================== 2. ARBOL DE PROBLEMAS =====================
h1("2. Árbol de Problemas")
para("Problema central: servicio de acueducto deficiente, discontinuo y con agua no apta en la cabecera de Tadó, "
     "pese a estar en una de las zonas más lluviosas del planeta. Los efectos se ubican en la parte superior y las "
     "causas en la parte inferior.")
add_image_centered("arbol_problemas.png")

# ===================== 3. ARBOL DE OBJETIVOS =====================
h1("3. Árbol de Objetivos")
para("Objetivo central: garantizar agua potable continua (24 h) y de calidad para la población de la cabecera de "
     "Tadó. Es el espejo positivo del árbol de problemas: cada causa se convierte en un medio y cada efecto en un fin.")
add_image_centered("arbol_objetivos.png")

# ===================== 4. MML =====================
h1("4. Matriz de Marco Lógico (MML)")
add_table(
    ["Nivel", "Resumen narrativo", "Indicadores", "Medios de verificación", "Supuestos"],
    [["FIN", "Contribuir al desarrollo humano y a la reducción de la pobreza en Tadó.",
      "IPM municipal; cobertura de necesidades básicas de agua.",
      "DANE (IPM), TerriData, Plan de Desarrollo.",
      "Estabilidad institucional y de orden público."],
     ["PROPÓSITO", "La población de la cabecera dispone de agua potable continua y de calidad.",
      "Continuidad ≥ 24 h; IRCA ≤ 5 (apta); cobertura ≥ 98 %.",
      "Reportes SUI/Superservicios; informes IRCA.",
      "La comunidad se conecta y paga la tarifa."],
     ["COMPONENTES", "1) PTAP 45 L/s. 2) Redes rehabilitadas y sectorizadas. 3) Micromedición.",
      "Capacidad instalada; km de red; IANC ≤ 25 %; N° de medidores.",
      "Actas de obra; catastro de redes; SUI.",
      "Suministro eléctrico/insumos estable; predios disponibles."],
     ["ACTIVIDADES", "Estudios y diseños; obras civiles; equipos; interventoría; gestión ambiental y social.",
      "Avance físico y financiero (%); cumplimiento del cronograma.",
      "Informes de interventoría; ejecución presupuestal.",
      "Recursos cofinanciados desembolsados a tiempo."]],
    col_align=["center", "left", "left", "left", "left"])

# ===================== 5. SOLUCION DE INGENIERIA =====================
h1("5. Solución de Ingeniería Propuesta")
h2("5.1 Diseño de caudal (Resolución 0330 de 2017 — RAS)")
para("La Resolución 0330 de 2017 (RAS) fija la metodología de diseño. Con población municipal de 17.000 hab "
     "(CNPV 2018, DANE), fracción urbana estimada del 60 % y crecimiento geométrico del 1,2 % anual, la población "
     "de diseño a 25 años es ≈ 14.764 hab. Para clima cálido (< 1.000 m s.n.m.) la dotación neta es 140 L/hab·día y "
     "las pérdidas máximas admisibles 25 %, de donde la dotación bruta es 186,7 L/hab·día.")
add_table(
    ["Caudal", "Fórmula", "Valor"],
    [["Medio diario (Qmd)", "Pob · Dot.bruta / 86.400", "31,9 L/s"],
     ["Máximo diario (QMD)", "Qmd · k1 (1,30)", "41,5 L/s"],
     ["Máximo horario (QMH)", "QMD · k2 (1,60)", "66,4 L/s"],
     ["Capacidad PTAP seleccionada", "—", "45 L/s"]],
    col_align=["left", "left", "right"])

h2("5.2 Esquema del sistema")
para("1. Captación (bocatoma) sobre el río San Juan, con manejo de crecientes y sólidos gruesos. "
     "2. Línea de aducción hasta la PTAP. "
     "3. PTAP convencional de 45 L/s con tren robusto frente a alta turbiedad: pretratamiento/desarenado → "
     "coagulación-floculación → sedimentación → filtración → desinfección; calidad verificada con el IRCA "
     "(Resolución 2115 de 2007). "
     "4. Almacenamiento (tanques) y estaciones de bombeo para presión y continuidad. "
     "5. Redes de distribución rehabilitadas y sectorizadas, con micromedición para reducir el IANC.")

h2("5.3 Presupuesto / CAPEX por capítulos")
para("El CAPEX total se ancla a la inversión real del proyecto (COP 19.971 MM). El desglose por capítulos es una "
     "estimación consistente con ese total y con el alcance del proyecto.")
add_table(
    ["Capítulo", "COP MM"],
    [["Captación (bocatoma) y línea de aducción", "2.400"],
     ["PTAP 45 L/s (obra civil + equipos)", "6.300"],
     ["Almacenamiento (tanques) y estaciones de bombeo", "2.800"],
     ["Redes de distribución, conexiones y micromedición", "4.200"],
     ["Optimización del alcantarillado (componente asociado)", "2.471"],
     ["Estudios, diseños, interventoría y gestión ambiental/social", "1.800"],
     ["TOTAL CAPEX", "19.971"]],
    col_align=["left", "right"])

# ===================== 6. ANALISIS FINANCIERO =====================
h1("6. Análisis Financiero")
h2("6.1 Estructura de financiación y WACC")
para("Los acueductos municipales pequeños no se financian con deuda pura: dependen de aportes públicos no "
     "reembolsables (SGR, PDA del Chocó, SGP, cooperación). Se adopta la estructura típica 80 % público / 10 % "
     "deuda / 10 % equity.")
add_table(
    ["Fuente", "%", "COP MM", "Costo"],
    [["Aporte público (SGR/PDA/MinVivienda/cooperación)", "80 %", "15.977", "tasa social 9 %"],
     ["Deuda (banca de desarrollo, FINDETER)", "10 %", "1.997", "Kd = 11 %"],
     ["Equity (operador/municipio)", "10 %", "1.997", "Ke = 14 %"]],
    col_align=["left", "center", "right", "center"])
para("WACC = 0,80·9 % + 0,10·14 % + 0,10·11 %·(1 − 0,35) = 9,31 %. El costo de oportunidad del aporte público se "
     "toma como la tasa social del DNP (9 %); la deuda incluye escudo fiscal con renta del 35 %.")

h2("6.2 Indicadores y evaluaciones")
add_table(
    ["Evaluación", "Tasa de descuento", "Resultado"],
    [["1) Financiera del proyecto (solo tarifas)", "WACC = 9,31 %", "VPN = −15.487 MM; TIR = −2,45 %"],
     ["2) Del inversionista (deuda + equity)", "Ke = 14 %", "VPN = −730 MM; TIR = 10,46 %"],
     ["3) Socioeconómica (tasa social DNP)", "9 %", "VPN = +19.626 MM; B/C = 1,56"]],
    col_align=["left", "center", "left"])
para("El proyecto no es rentable solo con tarifas (norma en acueductos del Pacífico). Es socialmente rentable: el "
     "VPN económico es ampliamente positivo y la relación beneficio/costo (1,56) supera 1, lo que justifica la "
     "inversión pública, coherente con el estándar internacional de la OMS.")

h2("6.3 Análisis de sensibilidad — aporte público vs. TIR del inversionista")
add_table(
    ["Aporte público", "TIR inversionista"],
    [["0 %", "−5,40 %"],
     ["40 %", "−1,46 %"],
     ["60 %", "+2,18 %"],
     ["70 %", "+5,20 %"],
     ["80 %", "+10,46 %"],
     ["90 %", "+25,18 %"]],
    col_align=["center", "right"])
para("Sin aporte público el proyecto es inviable para cualquier inversionista. Solo resulta atractivo para un "
     "operador privado con aporte público muy alto (≥ 85–90 %), o debe ser operado por una empresa "
     "pública/comunitaria que no exige el 14 % de rentabilidad. Esto explica por qué lo financió MinVivienda con "
     "cooperación internacional.")

h2("6.4 Flujos de caja (Proyecto, Inversionista, Banco y Social)")
flujos = leer_flujos()
if flujos:
    add_table(
        ["Año", "Ingresos", "OPEX", "EBITDA", "FC Proyecto", "FC Inversion.", "FC Banco", "Benef. social", "FC Social"],
        flujos,
        col_align=["center"] + ["right"] * 8, fsize=8)
else:
    para("[No se encontró resultados_modelo.csv. Ejecuta modelo_financiero.py para generarlo.]", italic=True, size=9)
para("FC Proyecto: rentabilidad de la inversión total (descontado al WACC). FC Inversionista: lo que recibe el "
     "accionista tras intereses y amortización (descontado al Ke). FC Banco: desembolsa la deuda en el año 0 y "
     "recibe la cuota fija 10 años. FC Social: beneficios sociales menos OPEX (descontado a la tasa social 9 %).",
     italic=True, size=9)

# ===================== 7. MATRIZ DE RIESGOS =====================
h1("7. Matriz de Riesgos (20 riesgos)")
para("Escala: Probabilidad (P) e Impacto (I) de 1 a 5. Severidad = P × I. "
     "Clasificación: Bajo (1–6), Medio (8–12), Alto (15–25).", italic=True, size=9)
riesgos = [
    ["1", "Crecientes/turbiedad extrema del río San Juan", "Ambiental", "4", "5", "20", "Alto"],
    ["2", "Contaminación por minería ilegal (mercurio)", "Ambiental", "4", "5", "20", "Alto"],
    ["3", "Inundaciones que dañan captación/redes", "Ambiental", "4", "4", "16", "Alto"],
    ["4", "Deforestación de la cuenca", "Ambiental", "3", "3", "9", "Medio"],
    ["5", "Vertimientos de aguas residuales a la fuente", "Ambiental", "3", "4", "12", "Medio"],
    ["6", "Resistencia comunitaria a micromedición/tarifa", "Social", "3", "3", "9", "Medio"],
    ["7", "Baja cultura de pago y cartera morosa", "Social", "4", "3", "12", "Medio"],
    ["8", "Presencia de grupos armados (extorsión/retrasos)", "Seguridad", "4", "4", "16", "Alto"],
    ["9", "Demora en desembolsos de cofinanciación", "Institucional", "4", "4", "16", "Alto"],
    ["10", "Debilidad del operador municipal", "Institucional", "3", "4", "12", "Medio"],
    ["11", "Sobrecostos por inflación de insumos/transporte", "Financiero", "3", "3", "9", "Medio"],
    ["12", "Tarifa insuficiente para cubrir OPEX", "Financiero", "4", "3", "12", "Medio"],
    ["13", "Diseño hidráulico subdimensionado", "Técnico", "2", "4", "8", "Medio"],
    ["14", "Fallas de equipos de bombeo", "Técnico", "3", "3", "9", "Medio"],
    ["15", "Suministro eléctrico inestable", "Técnico", "3", "4", "12", "Medio"],
    ["16", "Alto IANC por fugas no detectadas", "Operativo", "3", "3", "9", "Medio"],
    ["17", "Dificultad de acceso por mal estado de vías", "Logístico", "4", "3", "12", "Medio"],
    ["18", "Escasez de insumos químicos (coagulantes)", "Ambiental", "2", "3", "6", "Bajo"],
    ["19", "Demoras en predios/servidumbres", "Legal", "3", "3", "9", "Medio"],
    ["20", "Brote de EDA durante la transición de obra", "Salud pública", "2", "4", "8", "Medio"],
]
add_table(
    ["#", "Riesgo", "Tipo", "P", "I", "P×I", "Nivel"],
    riesgos,
    col_align=["center", "left", "left", "center", "center", "center", "center"])
para("Los riesgos más críticos (Altos) se concentran en factores ambientales (turbiedad, minería, inundaciones), "
     "seguridad e institucionales (desembolsos).")

# ===================== 8. INTERVENCION 10 RIESGOS =====================
h1("8. Intervención de 10 Riesgos (5 ambientales) y Recalificación")
para("Se intervienen los 10 riesgos prioritarios, incluidos 5 ambientales (riesgos 1, 2, 3, 5 y 4).")
inter = [
    ["1", "Turbiedad extrema del río San Juan", "Ambiental", "Pretratamiento robusto + sedimentadores de alta tasa + dosificación automática + tanque de regulación.", "8", "Medio"],
    ["2", "Contaminación por minería ilegal", "Ambiental", "Monitoreo de fuente + carbón activado + articulación con CODECHOCÓ y control de minería.", "9", "Medio"],
    ["3", "Inundaciones que dañan infraestructura", "Ambiental", "Captación elevada/protegida + protección de orillas + redes en cotas seguras.", "8", "Medio"],
    ["5", "Vertimientos a la fuente", "Ambiental", "Optimización de alcantarillado + campañas + coordinación municipal.", "6", "Bajo"],
    ["4", "Deforestación de la cuenca", "Ambiental", "Pago por servicios ambientales + reforestación de microcuenca con la comunidad.", "4", "Bajo"],
    ["8", "Grupos armados / inseguridad", "Seguridad", "Plan de seguridad de obra + articulación con autoridades + contratación local.", "8", "Medio"],
    ["9", "Demora en desembolsos", "Institucional", "Convenio con desembolsos por hitos + anticipo + fiducia.", "8", "Medio"],
    ["7", "Baja cultura de pago", "Social", "Subsidios focalizados + educación + tarifa social + facturación clara.", "6", "Bajo"],
    ["12", "Tarifa insuficiente", "Financiero", "Estudio tarifario CRA + subsidios cruzados + aporte SGP.", "6", "Bajo"],
    ["15", "Suministro eléctrico inestable", "Técnico", "Planta eléctrica de respaldo + almacenamiento con autonomía.", "6", "Bajo"],
]
add_table(
    ["#", "Riesgo", "Tipo", "Intervención propuesta", "P×I residual", "Nivel residual"],
    inter,
    col_align=["center", "left", "left", "left", "center", "center"])
para("Con las intervenciones, los riesgos pasan de niveles Alto/Medio a Medio/Bajo. Los 5 riesgos ambientales "
     "priorizados se reducen sustancialmente, lo cual es clave porque la fuente (río San Juan), la turbiedad extrema "
     "y la minería son las mayores amenazas para la operación.")

# ===================== 9. CONCLUSIONES =====================
h1("9. Conclusiones y Recomendaciones")
for x in [
    "Técnicamente viable: el caudal de diseño (QMD = 41,5 L/s, Res. 0330) y la PTAP de 45 L/s responden a la población a 25 años; el reto central es la alta turbiedad de la fuente.",
    "Financieramente NO rentable solo con tarifas (VPN = −15.487 MM; TIR = −2,45 %), lo cual es la norma en acueductos del Pacífico y no descalifica el proyecto.",
    "Requiere cofinanciación pública alta (≥ 80–90 %) para ejecutarse y operar de forma sostenible; lo natural es un operador público/comunitario apoyado por SGR/PDA/cooperación.",
    "Socialmente muy rentable: VPN social de +19.626 MM y B/C = 1,56 a la tasa social del DNP (9 %); aquí está la verdadera justificación de la inversión.",
    "Los riesgos ambientales y de seguridad son los más críticos; las medidas propuestas los reducen a niveles manejables.",
]:
    doc.add_paragraph(x, style="List Number")

# ===================== 10. FUENTES =====================
h1("10. Fuentes")
fuentes = [
    "DANE — Censo Nacional de Población y Vivienda 2018: población de Tadó ≈ 17.000 hab.",
    "Clima de Tadó (~7.921 mm/año, clima tropical lluvioso): Wikipedia — Tadó.",
    "Río San Juan (fuente hídrica, vertiente Pacífico): Wikipedia — San Juan River (Colombia).",
    "Proyecto real de acueducto de Tadó (MinVivienda + cooperación española; inversión COP 19.971 millones; meta 24 h y 98 % de cobertura): Agencia Anadolu (aa.com.tr).",
    "Contexto de seguridad en la región de Tadó (Chocó): El Ciudadano.",
    "Resolución 0330 de 2017 (RAS) — MinVivienda: dotación, pérdidas y coeficientes k1/k2.",
    "Resolución 2115 de 2007: criterios de calidad del agua (IRCA).",
    "Marco tarifario CRA (Resolución CRA 943 de 2021).",
    "Tasa social de descuento (9 %): Departamento Nacional de Planeación (DNP).",
]
for fnt in fuentes:
    doc.add_paragraph(fnt, style="List Bullet")
para("El contenido de las fuentes consultadas fue parafraseado y resumido para cumplir con las restricciones de "
     "licenciamiento.", italic=True, size=9)

doc.save("INFORME_PEPI.docx")
print("OK -> INFORME_PEPI.docx generado")
