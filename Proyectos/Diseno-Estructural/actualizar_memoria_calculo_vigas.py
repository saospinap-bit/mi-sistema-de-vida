#!/usr/bin/env python3
"""Actualiza la memoria estructural con el diseño reconciliado de 592 vigas.

Conserva el DOCX fuente, sustituye los apartados y tablas obsoletos de vigas,
y genera una copia final trazable con el libro canónico y las láminas B2.
No cambia armados ni convierte estados NO CUMPLE en resultados favorables.
"""
from __future__ import annotations

from collections import Counter, defaultdict
from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from openpyxl import load_workbook

BASE = Path(__file__).resolve().parent
SOURCE = BASE / "GRUPO_6_SANTA_MARTA_MEMORIA_ACTUALIZADA (6).docx"
DESIGN = BASE / "Diseño de vigas proyecto diseño DEF.xlsx"
PREVIEW = BASE / "Planos-Autocad-B2-Vigas-Continuas" / "VISTA-PREVIA-TODAS-LAS-LAMINAS.png"
OUTPUT = BASE / "GRUPO_6_SANTA_MARTA_MEMORIA_CALCULO_VIGAS_ACTUALIZADA.docx"
DETAIL_SHEET = "Todas las Vigas (592)"
SUMMARY_SHEETS = ("Vigas de Carga (7)", "Vigas de Rigidez (5)")
EXPECTED_FORMULAS = 21_744
EXPECTED_FRAMES = 592
EXPECTED_GROUPS = 12

NAVY = "1F4E78"
LIGHT_BLUE = "D9EAF7"
GREEN = "C6E0B4"
RED = "F4CCCC"
GRAY = "E7E6E6"
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)


def clean(value) -> str:
    return " ".join(str(value).split())


def fmt(value, decimals=1) -> str:
    return f"{float(value):.{decimals}f}"


def validate_ooxml_book() -> int:
    with ZipFile(DESIGN) as archive:
        names = archive.namelist()
        external = [name for name in names if "externalLinks" in name]
        if external:
            raise ValueError(f"Persisten vínculos externos OOXML: {external}")
        worksheets = [
            archive.read(name)
            for name in names
            if name.startswith("xl/worksheets/sheet") and name.endswith(".xml")
        ]
    formulas = sum(xml.count(b"<f") for xml in worksheets)
    payload = b"".join(worksheets)
    if formulas != EXPECTED_FORMULAS:
        raise ValueError(f"Fórmulas={formulas}; esperadas={EXPECTED_FORMULAS}")
    if b"[1]" in payload:
        raise ValueError("Persisten referencias [1] en el libro")
    if b't="e"' in payload:
        raise ValueError("El libro contiene errores Excel almacenados")
    return formulas


def load_design_data():
    formulas = validate_ooxml_book()
    workbook = load_workbook(DESIGN, read_only=True, data_only=True, keep_links=False)
    if workbook.sheetnames[:4] != ["Parametros", *SUMMARY_SHEETS, DETAIL_SHEET]:
        raise ValueError(f"Hojas inesperadas: {workbook.sheetnames}")

    parameters = [workbook["Parametros"].cell(row, 2).value for row in range(2, 13)]
    if parameters != [28, 420, 420, 0.9, 0.75, 40, 199, 9.5, 25, None, 15.9]:
        raise ValueError(f"Parámetros internos inesperados: {parameters}")

    groups = {}
    for sheet_name in SUMMARY_SHEETS:
        sheet = workbook[sheet_name]
        headers = [cell.value for cell in sheet[3]]
        for values in sheet.iter_rows(min_row=4, values_only=True):
            if values[0] is None:
                continue
            row = dict(zip(headers, values))
            groups[str(row["Grupo"])] = row

    detail = workbook[DETAIL_SHEET]
    headers = [cell.value for cell in detail[3]]
    detail_status = defaultdict(Counter)
    frames = set()
    excel_errors = []
    for values in detail.iter_rows(min_row=4, values_only=True):
        if values[0] is None:
            continue
        row = dict(zip(headers, values))
        frame = str(row["Frame"])
        if frame in frames:
            raise ValueError(f"Frame repetido en la hoja detallada: {frame}")
        frames.add(frame)
        detail_status[str(row["Grupo"])][str(row["ESTADO"])] += 1
        for header, value in row.items():
            if isinstance(value, str) and value.startswith("#"):
                excel_errors.append(f"{frame}:{header}={value}")
    workbook.close()

    if len(frames) != EXPECTED_FRAMES:
        raise ValueError(f"Frames detallados={len(frames)}; esperados={EXPECTED_FRAMES}")
    if len(groups) != EXPECTED_GROUPS:
        raise ValueError(f"Grupos resumen={len(groups)}; esperados={EXPECTED_GROUPS}")
    if excel_errors:
        raise ValueError(f"Errores almacenados: {excel_errors[:10]}")
    if sum(sum(counter.values()) for counter in detail_status.values()) != EXPECTED_FRAMES:
        raise ValueError("Los estados detallados no cubren los 592 frames")
    if any(not set(counter).issubset({"CUMPLE", "NO CUMPLE"}) for counter in detail_status.values()):
        raise ValueError(f"Estados detallados inesperados: {dict(detail_status)}")

    return groups, detail_status, formulas


def find_paragraph(document: Document, text: str):
    for index, paragraph in enumerate(document.paragraphs):
        if clean(paragraph.text) == text:
            return index
    raise ValueError(f"No se encontró el párrafo: {text}")


def set_text(paragraph, text: str, bold_prefix: str | None = None):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        first.bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)


def set_cell_shading(cell, fill: str):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=40, start=45, bottom=40, end=45):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def style_cell(cell, header=False, fill: str | None = None, font_size=7.0):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    if fill:
        set_cell_shading(cell, fill)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
            run.bold = header
            run.font.color.rgb = WHITE if header else BLACK


def replace_table(table, headers, rows, widths=None, status_column=None):
    while len(table.columns) < len(headers):
        table.add_column(Cm(1.5))
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    header = table.rows[0]
    for index, value in enumerate(headers):
        header.cells[index].text = str(value)
        style_cell(header.cells[index], header=True, fill=NAVY, font_size=7.0)
    set_repeat_table_header(header)

    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.text = str(value)
            fill = None
            if status_column is not None and index == status_column:
                fill = GREEN if "NO CUMPLE" not in str(value) else RED
            style_cell(cell, fill=fill, font_size=6.8)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)


def failure_causes(row) -> str:
    checks = (
        ("Flexión −", "flexión −"),
        ("Flexión +", "flexión +"),
        ("Sección cortante", "límite de Vs"),
        ("Cortante", "cortante"),
        ("Chequeo torsión", "torsión"),
        ("Chequeo V-T", "interacción V-T"),
        ("Chequeo Al", "Al torsional"),
    )
    failed = []
    for column, label in checks:
        value = str(row[column])
        if value not in {"CUMPLE", "DESPRECIABLE"}:
            failed.append(label)
    return ", ".join(failed) if failed else "—"


def group_table_rows(groups, detail_status, family: str):
    names = (
        ["VC1", "VC2", "VC3", "VC4", "VC5", "VC6", "VC7"]
        if family == "carga"
        else ["VR1", "VR N1", "VR2", "VR3", "VRAUX"]
    )
    rows = []
    for name in names:
        row = groups[name]
        status = str(row["ESTADO"])
        failed = detail_status[name]["NO CUMPLE"]
        total = sum(detail_status[name].values())
        status_text = f"{status}\nDetalle: {failed}/{total} NO CUMPLE"
        causes = failure_causes(row)
        if causes != "—":
            status_text += f"\nRev.: {causes}"
        rows.append((
            name,
            int(row["N vigas"]),
            f"{fmt(row['Ln adopt. mín (m)'], 2)}–{fmt(row['Ln adopt. máx (m)'], 2)}",
            f"{fmt(row['Mu− ENVFLEX (kN·m)'])} / {fmt(row['Mu+ ENVFLEX (kN·m)'])}",
            f"{fmt(row['Vu ENVFLEX (kN)'])} / {fmt(row['Vu ENVCORT (kN)'])}",
            fmt(row["Tu diseño compat. (kN·m)"]),
            f"{int(row['Nº5 sup'])}#5 / {int(row['Nº5 inf'])}#5 / {int(row['Nº5 torsión dedicadas'])}#5",
            f"{clean(row['Estribo'])}; {int(row['Ramas'])}R\n@{int(row['s extremo DMO (mm; zona 2h)'])}/@{int(row['s centro (mm)'])} mm\n{status_text}",
        ))
    return rows


def set_update_fields(document: Document):
    settings = document.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def add_traceability_annex(document: Document, groups, detail_status, formulas):
    document.add_page_break()
    heading = document.add_heading("ANEXO — TRAZABILIDAD DEL DISEÑO DE VIGAS", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    total_status = Counter()
    for counter in detail_status.values():
        total_status.update(counter)
    formula_text = f"{formulas:,}".replace(",", ".")
    paragraph = document.add_paragraph()
    set_text(
        paragraph,
        "La entrega se reconstruyó a partir de resultados sap.xlsx y geomatria sap.xlsx. "
        f"El libro Diseño de vigas proyecto diseño DEF.xlsx conserva {formula_text} fórmulas, "
        "contiene parámetros internos, no presenta vínculos [1] ni errores Excel almacenados, "
        f"y cubre {EXPECTED_FRAMES} frames distribuidos en {EXPECTED_GROUPS} grupos. "
        f"El recálculo produjo {total_status['CUMPLE']} estados CUMPLE y "
        f"{total_status['NO CUMPLE']} estados NO CUMPLE en la hoja detallada.",
    )
    paragraph = document.add_paragraph()
    set_text(
        paragraph,
        "Los estados anteriores son verificaciones automáticas del libro; no constituyen "
        "aprobación estructural. Los grupos o frames NO CUMPLE deben corregirse, recalcularse "
        "y someterse a revisión profesional antes de cualquier emisión para construcción.",
    )
    document.add_heading("Archivos que respaldan este capítulo", level=2)
    for item in (
        "Diseño de vigas proyecto diseño DEF.xlsx — cálculo por grupo y por frame.",
        "resultados sap.xlsx — fuerzas ENVFLEX y ENVCORT y asignaciones de grupos.",
        "geomatria sap.xlsx — conectividad, luces y secciones del modelo.",
        "Planos-Autocad-B2-Vigas-Continuas/ — DXF, PDF, PNG y paquete ZIP.",
        "reconciliar_diseno_vigas.py — reconciliación y recálculo reproducible.",
        "generar_planos_vigas_continuas_b2.py — generación y auditoría CAD.",
    ):
        bullet = document.add_paragraph()
        set_text(bullet, f"• {item}")

    if PREVIEW.exists():
        document.add_heading("Vista previa de las láminas B2 de vigas", level=2)
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(PREVIEW), width=Inches(6.45))
        caption = document.add_paragraph(
            "Ilustración. Láminas B2 reconciliadas de vigas de carga y vigas de rigidez."
        )
        caption.style = "Caption"
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note = document.add_paragraph()
        set_text(note, "DETALLE ACADÉMICO — NO EMITIR PARA CONSTRUCCIÓN")
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in note.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(192, 0, 0)


def main():
    groups, detail_status, formulas = load_design_data()
    document = Document(SOURCE)
    set_update_fields(document)

    # Portada y fecha de la versión actualizada.
    document.paragraphs[1].text = "18/07/2026"
    document.paragraphs[7].text = "CONTENIDO: MEMORIA ESTRUCTURAL ACTUALIZADA — DISEÑO DE VIGAS"
    document.paragraphs[31].text = "FECHA DE ACTUALIZACIÓN: 18/07/2026"

    load_index = find_paragraph(document, "Diseño de las vigas de carga")
    set_text(
        document.paragraphs[load_index + 1],
        "Las vigas de carga se reconciliaron directamente con los exportes vigentes del modelo "
        "SAP2000. Se verificaron 240 frames agrupados en VC1 a VC7. Todas las secciones del "
        "modelo y adoptadas son de 450 × 550 mm; las luces individuales varían entre 2.65 y "
        "5.10 m. Para cada frame se extrajeron los máximos de momento negativo y positivo y "
        "cortante de ENVFLEX, así como el máximo cortante y la torsión de ENVCORT. El resumen "
        "de cada grupo adopta la envolvente más desfavorable sin mezclar identificadores de frame.",
    )
    set_text(
        document.paragraphs[load_index + 2],
        "Diseño a flexión. Con f'c = 28 MPa, fy = 420 MPa, φ = 0.90, recubrimiento de 40 mm "
        "y barras longitudinales #5, el peralte efectivo se calcula como d = h − cc − db,est − "
        "db,long/2. Se determina As requerido con Rn y ρ, se adopta como mínimo el mayor entre "
        "(1.4/fy)b·d y (0.25√f'c/fy)b·d, y se verifica φMn = φ As fy(d − a/2) ≥ Mu, con "
        "a = As fy/(0.85 f'c b). Las barras provistas y las capas se comprueban por separado en "
        "las caras superior e inferior.",
        "Diseño a flexión.",
    )
    set_text(
        document.paragraphs[load_index + 3],
        "Cortante, capacidad y torsión. Se calcula Vc = 0.17√f'c·b·d y el cortante de diseño "
        "como el mayor entre Vu de ENVCORT y Vcap, donde Vcap incorpora los momentos probables "
        "Mpr obtenidos con 1.25fy en los extremos. El refuerzo transversal satisface la demanda, "
        "los límites de Vs y las separaciones DMO. Para torsión se evalúa el umbral φTth = "
        "φ·0.083√f'c·Acp²/pcp; cuando Tu lo supera se calculan At/s y el acero longitudinal Al, "
        "además de la interacción combinada V-T. El diámetro real del estribo se toma de cada "
        "fila: #3 (9.5 mm) o #4 (12.7 mm), evitando aplicar un único diámetro a todos los grupos.",
        "Cortante, capacidad y torsión.",
    )
    vc1 = groups["VC1"]
    set_text(
        document.paragraphs[load_index + 4],
        "Ejemplo reconciliado VC1. Para b × h = 450 × 550 mm y d = "
        f"{fmt(vc1['d sup (mm)'], 2)} mm, Mu− = {fmt(vc1['Mu− ENVFLEX (kN·m)'], 2)} kN·m "
        f"exige As = {fmt(vc1['As req. sup (mm²)'], 1)} mm²; se proporcionan "
        f"{int(vc1['Nº5 sup'])}#5 = {fmt(vc1['As prov. sup (mm²)'], 0)} mm² y "
        f"φMn− = {fmt(vc1['φMn− (kN·m)'], 1)} kN·m. Para Mu+ = "
        f"{fmt(vc1['Mu+ ENVFLEX (kN·m)'], 2)} kN·m se disponen "
        f"{int(vc1['Nº5 inf'])}#5 y φMn+ = {fmt(vc1['φMn+ (kN·m)'], 1)} kN·m. "
        f"El cortante de diseño es {fmt(vc1['Vu diseño (kN)'], 1)} kN y se atiende con "
        f"{clean(vc1['Estribo'])} de {int(vc1['Ramas'])} ramas @"
        f"{int(vc1['s extremo DMO (mm; zona 2h)'])}/@{int(vc1['s centro (mm)'])} mm. "
        f"Como Tu = {fmt(vc1['Tu diseño compat. (kN·m)'], 2)} kN·m supera φTth = "
        f"{fmt(vc1['φTth (kN·m)'], 2)} kN·m, se incorporan "
        f"{int(vc1['Nº5 torsión dedicadas'])}#5 dedicadas a torsión. VC1 resulta CUMPLE.",
        "Ejemplo reconciliado VC1.",
    )
    document.paragraphs[load_index + 5].text = (
        "Tabla 61. Resumen reconciliado de vigas de carga (450 × 550 mm)."
    )

    stiffness_index = find_paragraph(document, "Diseño de las riostras o vigas de rigidez")
    document.paragraphs[stiffness_index].text = "Diseño de las vigas de rigidez"
    set_text(
        document.paragraphs[stiffness_index + 1],
        "Las vigas de rigidez comprenden 352 frames agrupados en VR1, VR N1, VR2, VR3 y "
        "VRAUX. VR1, VR N1, VR2 y VR3 emplean secciones de 500 × 550 mm; VRAUX utiliza "
        "450 × 550 mm. Las luces varían entre 1.50 y 5.50 m. Se aplicó el mismo procedimiento "
        "de flexión, cortante por capacidad, torsión e interacción descrito para las vigas de "
        "carga, manteniendo la trazabilidad entre grupo, frame gobernante y envolvente SAP.",
    )
    set_text(
        document.paragraphs[stiffness_index + 2],
        "Resultados y condición de revisión. En la hoja detallada se obtuvieron 538 frames "
        "CUMPLE y 54 frames NO CUMPLE. En el resumen grupal únicamente VC1, VR1 y VRAUX "
        "alcanzan CUMPLE con todas las comprobaciones actuales; los nueve grupos restantes "
        "conservan NO CUMPLE por una o más verificaciones de flexión, cortante, torsión, acero "
        "longitudinal torsional o interacción V-T. Estos resultados no se corrigieron alterando "
        "armados de forma automática: deben revisarse y rediseñarse antes de una eventual "
        "emisión profesional. La Tabla 62 identifica además cuántos frames detallados fallan "
        "dentro de cada grupo.",
        "Resultados y condición de revisión.",
    )
    document.paragraphs[stiffness_index + 3].text = (
        "Tabla 62. Resumen reconciliado de vigas de rigidez (500 × 550 mm; VRAUX 450 × 550 mm)."
    )

    headers = (
        "Grupo", "N", "L mín–máx\n(m)", "Mu− / Mu+\n(kN·m)",
        "Vu F / Vu C\n(kN)", "Tu\n(kN·m)", "Sup / Inf / Tor", "Estribo y estado",
    )
    widths = (1.15, 0.7, 1.25, 1.65, 1.65, 1.1, 2.05, 3.5)
    replace_table(
        document.tables[10], headers, group_table_rows(groups, detail_status, "carga"),
        widths=widths, status_column=7,
    )
    replace_table(
        document.tables[11], headers, group_table_rows(groups, detail_status, "rigidez"),
        widths=widths, status_column=7,
    )

    # Actualiza detallado y longitudes para las secciones vigentes.
    development_index = find_paragraph(document, "Longitudes de desarrollo, ganchos y estribos")
    set_text(
        document.paragraphs[development_index + 1],
        "Para el despiece se adoptan las longitudes verificadas en la entrega parcial y en el "
        "generador CAD: ld inferior = 370 mm, ld superior = 480 mm y ldh para barra #5 = "
        "305 mm. Los ganchos de las barras principales son de 90°, con cola de 195 mm y "
        "diámetro interior de doblado de 100 mm. Los estribos cerrados terminan en ganchos "
        "sísmicos de 135°. La aplicabilidad y disponibilidad geométrica de cada anclaje debe "
        "confirmarse en la revisión final de nudos.",
    )
    document.paragraphs[development_index + 2].text = (
        "Tabla 64. Longitudes de desarrollo y geometría de ganchos adoptadas (mm)."
    )
    set_text(
        document.paragraphs[development_index + 3],
        "Las barras superiores de apoyo se prolongan como mínimo un cuarto de las luces "
        "adyacentes más el desarrollo requerido; las inferiores se detallan por tramo con "
        "anclaje en ambos extremos. Las longitudes definitivas se obtienen frame por frame a "
        "partir de la geometría SAP y no mediante una única luz representativa para toda la familia.",
    )
    document.paragraphs[development_index + 4].text = (
        "Tabla 65. Criterio de longitud de corte de barras longitudinales."
    )
    set_text(
        document.paragraphs[development_index + 5],
        "La zona de confinamiento se extiende 2h = 1 100 mm desde la cara de cada apoyo. "
        "Las separaciones calculadas varían por grupo y se reproducen en las Tablas 61 y 62; "
        "no debe sustituirse esa información por una separación única. Se emplean estribos "
        "cerrados #3 o #4, de dos o cuatro ramas, según demanda de cortante y torsión.",
    )
    document.paragraphs[development_index + 6].text = (
        "Tabla 66. Separaciones de estribos y zonas de confinamiento vigentes."
    )

    replace_table(
        document.tables[13],
        ("Barra", "db", "ldh", "Cola 90°", "Gancho 135°", "Doblado"),
        (
            ("#3", "9.5", "181", "114", "75", "57"),
            ("#4", "12.7", "242", "152", "76", "76"),
            ("#5", "15.9", "305", "195", "95", "100"),
        ),
        widths=(1.4, 1.2, 1.2, 1.5, 1.7, 1.4),
    )
    replace_table(
        document.tables[14],
        ("Elemento", "Barra", "Rango de luz (m)", "Anclaje", "Criterio de corte"),
        (
            ("Viga de carga", "#5", "2.65–5.10", "ldh 305 mm", "Por frame: Lclara + anclajes"),
            ("Viga de rigidez", "#5", "1.50–5.50", "ldh 305 mm", "Por frame: Lclara + anclajes"),
            ("Vigueta", "#4", "hasta 4.80", "ldh 242 mm", "Según entrega parcial"),
        ),
        widths=(2.4, 1.2, 2.0, 2.0, 3.1),
    )
    replace_table(
        document.tables[15],
        ("Elemento", "d (mm)", "Zona 2h (mm)", "s extremo (mm)", "s centro (mm)", "Configuración"),
        (
            ("Vigas de carga", "489.35–492.55", "1 100", "70–120", "70–200", "#3/#4; 2/4 ramas"),
            ("Vigas de rigidez", "489.35–492.55", "1 100", "40–120", "40–180", "#3/#4; 2/4 ramas"),
            ("Vigueta", "304", "680", "Según diseño", "Según diseño", "Entrega parcial"),
        ),
        widths=(2.4, 2.0, 1.8, 2.0, 2.0, 2.5),
    )

    despiece_index = find_paragraph(document, "Despiece de los elementos")
    set_text(
        document.paragraphs[despiece_index + 1],
        "El despiece de vigas se presenta en cuatro láminas B2 agrupadas por familia. Cada "
        "tramo dibujado identifica el frame SAP y rotula Mu−, Mu+, Vu de ENVFLEX, Vu de "
        "ENVCORT, Tu y su estado calculado. Los paneles laterales muestran sección, refuerzo, "
        "estribos, estado de grupo y cantidad de frames NO CUMPLE. Los planos conservan la "
        "leyenda DETALLE ACADÉMICO — NO EMITIR PARA CONSTRUCCIÓN y no constituyen planos "
        "aprobados ni base autorizada para armado en obra.",
    )

    annex_index = find_paragraph(document, "ANEXOS")
    if annex_index + 3 < len(document.paragraphs):
        document.paragraphs[annex_index + 2].text = "Libro reconciliado: Diseño de vigas proyecto diseño DEF.xlsx"
        document.paragraphs[annex_index + 3].text = "Planos: Planos-Autocad-B2-Vigas-Continuas/"

    add_traceability_annex(document, groups, detail_status, formulas)
    document.core_properties.title = "Memoria estructural actualizada — diseño reconciliado de vigas"
    document.core_properties.subject = "Diseño de 592 vigas en 12 grupos con trazabilidad SAP"
    document.core_properties.comments = (
        "Entrega académica. Conserva estados NO CUMPLE y no autoriza construcción."
    )
    document.save(OUTPUT)
    print(f"OK {OUTPUT}")
    print(f"Grupos: {len(groups)} | frames: {EXPECTED_FRAMES} | fórmulas: {formulas}")
    print("ADVERTENCIA: los estados NO CUMPLE se conservan para revisión profesional.")


if __name__ == "__main__":
    main()
