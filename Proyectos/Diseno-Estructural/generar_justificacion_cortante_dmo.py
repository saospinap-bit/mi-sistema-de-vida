#!/usr/bin/env python3
"""Genera la justificación y actualiza la memoria con el diseño final de vigas."""
from __future__ import annotations

import importlib.util
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import load_workbook

BASE = Path(__file__).resolve().parent
DESIGN_SCRIPT = BASE / "generar_diseno_vigas_envolventes.py"
SOURCE_MEMORY = BASE / "GRUPO_6_SANTA_MARTA_MEMORIA_ACTUALIZADA (6).docx"
EXCEL_FINAL = BASE / "DISENO-VIGAS-ENTREGA-FINAL-DMO.xlsx"
OUTPUT = BASE / "JUSTIFICACION-DISENO-CORTANTE-DMO.docx"
MEMORY_OUTPUT = BASE / "GRUPO_6_SANTA_MARTA_MEMORIA_FINAL_VIGAS.docx"


def load_design_module():
    spec = importlib.util.spec_from_file_location("beam_design", DESIGN_SCRIPT)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False, color=None, size=7):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_equation(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Cambria Math"
    run.font.size = Pt(11)


def configure(doc):
    section = doc.sections[-1]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "808080")
        borders.append(element)
    tbl_pr.append(borders)


GROUP_SHEETS = ("Vigas de Carga (7)", "Vigas de Rigidez (5)")
SUMMARY_HEADERS = [
    "Grupo", "Frame", "N vigas", "Sección", "Ln mín–máx (m)",
    "Mu− (kN·m)", "Mu+ (kN·m)", "Vu ENVCORT (kN)",
    "Vu diseño (kN)", "Tu diseño (kN·m)", "Estribo", "Estado",
]
DETAIL_HEADERS = [
    "Grupo", "Nº5 sup.", "Nº5 inf.", "Nº5 torsión", "Ramas",
    "s extremo (mm)", "s centro (mm)", "φVn (kN)",
    "Interacción V–T", "Límite V–T", "Estado",
]


def read_excel_group_rows():
    """Lee literalmente los resúmenes con valores cacheados del Excel final."""
    if not EXCEL_FINAL.exists():
        raise FileNotFoundError(f"No existe el Excel final: {EXCEL_FINAL}")
    workbook = load_workbook(EXCEL_FINAL, data_only=True, read_only=True)
    rows = []
    for sheet_name in GROUP_SHEETS:
        sheet = workbook[sheet_name]
        headers = {cell.value: cell.column for cell in sheet[3] if cell.value}
        required = {
            "Grupo", "Frame", "N vigas", "Ln adopt. mín (m)", "Ln adopt. máx (m)",
            "b adopt. (mm)", "h adopt. (mm)", "Mu− ENVFLEX (kN·m)",
            "Mu+ ENVFLEX (kN·m)", "Vu ENVCORT (kN)", "Tu diseño compat. (kN·m)",
            "Nº5 sup", "Nº5 inf", "Vu diseño (kN)", "Estribo", "Ramas",
            "s extremo DMO (mm; zona 2h)", "s centro (mm)", "φVn (kN)",
            "Nº5 torsión dedicadas", "ESTADO", "Interacción V-T", "Límite V-T",
        }
        missing = required - set(headers)
        if missing:
            raise ValueError(f"Faltan columnas en {sheet_name}: {sorted(missing)}")
        for row_number in range(4, sheet.max_row + 1):
            group = sheet.cell(row_number, headers["Grupo"]).value
            if not group:
                continue
            row = {header: sheet.cell(row_number, column).value
                   for header, column in headers.items()}
            row["_sheet"] = sheet_name
            rows.append(row)
    workbook.close()
    if len(rows) != 12:
        raise ValueError(f"Se esperaban 12 grupos en el Excel final y se leyeron {len(rows)}")
    return rows


def fmt_number(value, decimals=4):
    if value is None:
        return "—"
    text = f"{float(value):.{decimals}f}"
    return text.rstrip("0").rstrip(",").rstrip(".").replace(".", ",")


def validate_excel_rows(design, groups, excel_rows):
    """Impide emitir una memoria si el cálculo y el Excel final divergen."""
    by_group = {row["Grupo"]: row for row in excel_rows}
    expected = set(design.GROUP_ORDER)
    if set(by_group) != expected:
        raise ValueError(f"Grupos del Excel distintos a los esperados: {set(by_group) ^ expected}")
    numeric_pairs = [
        ("count", "N vigas"), ("b", "b adopt. (mm)"), ("h", "h adopt. (mm)"),
        ("Mu_neg", "Mu− ENVFLEX (kN·m)"), ("Mu_pos", "Mu+ ENVFLEX (kN·m)"),
        ("Vdesign", "Vu diseño (kN)"), ("Tu", "Tu diseño compat. (kN·m)"),
    ]
    for group_result in groups:
        excel = by_group[group_result["group"]]
        if str(excel["Frame"]) != str(group_result["frame"]):
            raise ValueError(f"Frame gobernante distinto en {group_result['group']}")
        for result_key, excel_key in numeric_pairs:
            if abs(float(group_result[result_key]) - float(excel[excel_key])) > 1e-6:
                raise ValueError(
                    f"{group_result['group']} {excel_key}: cálculo={group_result[result_key]} "
                    f"Excel={excel[excel_key]}"
                )
        expected_state = "CUMPLE" if group_result["group_all_ok"] else "NO CUMPLE"
        if excel["ESTADO"] != expected_state:
            raise ValueError(
                f"Estado distinto en {group_result['group']}: cálculo={expected_state}, "
                f"Excel={excel['ESTADO']}"
            )


def _styled_table(doc, headers, rows, design, font_size=5.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, bold=True, color="FFFFFF", size=font_size)
        shade(table.rows[0].cells[i], "17365D")
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            set_cell(cells[i], value, size=font_size)
        group = str(values[0])
        shade(cells[0], design.COLORS[group])
        shade(cells[-1], "C6E0B4" if str(values[-1]) == "CUMPLE" else "F4CCCC")
    return table


def excel_summary_table(doc, design, excel_rows):
    values = []
    for row in excel_rows:
        values.append([
            row["Grupo"], row["Frame"], row["N vigas"],
            f"{int(row['b adopt. (mm)'])}×{int(row['h adopt. (mm)'])}",
            f"{fmt_number(row['Ln adopt. mín (m)'], 5)}–{fmt_number(row['Ln adopt. máx (m)'], 5)}",
            fmt_number(row["Mu− ENVFLEX (kN·m)"], 4),
            fmt_number(row["Mu+ ENVFLEX (kN·m)"], 4),
            fmt_number(row["Vu ENVCORT (kN)"], 4),
            fmt_number(row["Vu diseño (kN)"], 6),
            fmt_number(row["Tu diseño compat. (kN·m)"], 4),
            row["Estribo"], row["ESTADO"],
        ])
    return _styled_table(doc, SUMMARY_HEADERS, values, design, font_size=5.2)


def excel_detail_table(doc, design, excel_rows):
    values = []
    for row in excel_rows:
        interaction = fmt_number(row["Interacción V-T"], 5)
        limit = fmt_number(row["Límite V-T"], 5)
        values.append([
            row["Grupo"], row["Nº5 sup"], row["Nº5 inf"],
            row["Nº5 torsión dedicadas"], row["Ramas"],
            fmt_number(row["s extremo DMO (mm; zona 2h)"], 0),
            fmt_number(row["s centro (mm)"], 0),
            fmt_number(row["φVn (kN)"], 4), interaction, limit, row["ESTADO"],
        ])
    return _styled_table(doc, DETAIL_HEADERS, values, design, font_size=5.6)


def append_viga_chapter(doc, design, records, results, groups, excel_rows):
    doc.add_page_break()
    add_heading(doc, "VERIFICACIÓN Y DESPIECE DE VIGAS — NO EMITIR PARA CONSTRUCCIÓN", 1)
    doc.add_paragraph(
        "Este capítulo verifica las 592 vigas con el export resultados sap.xlsx del modelo reanalizado en "
        "SAP2000 v26. Se conservaron los 12 grupos y sus secciones efectivamente asignadas: 450×550 mm "
        "para VC1–VC7 y VRAUX, y 500×550 mm para VR1, VR N1, VR2 y VR3. La verificación es conservadora "
        "y mantiene visibles todos los incumplimientos; no constituye liberación para construcción."
    )

    add_heading(doc, "Solicitaciones y criterios", 2)
    criteria = [
        "Mu− y Mu+ se obtienen de los extremos mínimo y máximo de M3 en ENVFLEX.",
        "Vu de diseño es el mayor entre ENVCORT directo y Ve por capacidad DMO.",
        "Ve = Vu,ENVFLEX + (Mpr− + Mpr+)/Ln, con Ln = máx(L estación − 0.60 m; 0.50L).",
        "R = R0·φa·φp·φr = 5·1·1·1 = 5; las fuerzas exportadas no se vuelven a dividir por R.",
        "Tu de diseño es el mayor |T| de ENVFLEX y ENVCORT. No se limita automáticamente a φTcr porque no se dispone de clasificación frame a frame ni de reanálisis de redistribución.",
        "El refuerzo transversal satisface máx(Av/s por resistencia; Av/s mínimo) + 2At/s, interacción V–T y límites de separación.",
        "Al se cubre con barras Nº5 dedicadas y distribuidas alrededor del perímetro, adicionales al acero requerido por flexión; no se acredita doblemente As superior/inferior.",
        "La zona extrema de estribos tiene longitud 2h desde cada cara y el primer estribo se coloca a máximo 50 mm.",
    ]
    for text in criteria:
        doc.add_paragraph("• " + text)

    add_equation(doc, "Vu,diseño = máx[Vu,ENVCORT ; Vu,ENVFLEX + (Mpr− + Mpr+)/Ln]")
    add_equation(doc, "Tu,diseño = máx(|TENVFLEX|, |TENVCORT|)  [sin reducción automática a φTcr]")

    add_heading(doc, "Resultados por grupo — valores del Excel final", 2)
    doc.add_paragraph(
        "Las siguientes tablas se leen directamente de DISENO-VIGAS-ENTREGA-FINAL-DMO.xlsx. "
        "Los valores corresponden a las hojas Vigas de Carga (7) y Vigas de Rigidez (5); "
        "no se conservan tablas de versiones preliminares."
    )
    vc_rows = [row for row in excel_rows if row["_sheet"] == "Vigas de Carga (7)"]
    vr_rows = [row for row in excel_rows if row["_sheet"] == "Vigas de Rigidez (5)"]
    add_heading(doc, "Vigas de carga VC1–VC7", 3)
    excel_summary_table(doc, design, vc_rows)
    add_heading(doc, "Vigas de rigidez VR1, VR N1, VR2, VR3 y VRAUX", 3)
    excel_summary_table(doc, design, vr_rows)

    add_heading(doc, "Armado y estribos del frame gobernante por grupo", 2)
    doc.add_paragraph(
        "El armado y los controles indicados son los de la fila gobernante de cada grupo en el Excel final. "
        "El estado de grupo exige además que todos sus frames individuales satisfagan los controles."
    )
    excel_detail_table(doc, design, vc_rows)
    excel_detail_table(doc, design, vr_rows)
    doc.add_paragraph(
        f"Resultado individual conservador: {sum(x['overall'] for x in results)} de {len(results)} vigas cumplen "
        f"todos los controles con torsión bruta. {sum(not x['overall'] for x in results)} vigas quedan NO CUMPLE. "
        f"En {sum(x['requires_compatibility_assessment'] for x in results)} vigas Tu bruto excede la referencia φTcr; "
        "cualquier reducción futura exige demostrar compatibilidad y reanalizar los elementos receptores."
    )

    failures = [x for x in results if not x["overall"]]
    add_heading(doc, "Observaciones obligatorias", 2)
    if failures:
        failure_counts = Counter(x["group"] for x in failures)
        summary = ", ".join(f"{group}: {failure_counts[group]}" for group in design.GROUP_ORDER if failure_counts[group])
        doc.add_paragraph(
            f"Quedan {len(failures)} frames no conformes ({summary}). Las causas son exceso de interacción V–T, "
            "límite seccional o ausencia de un estribo práctico con la sección actual. Deben aumentarse secciones, "
            "actualizarse las propiedades en SAP, reanalizarse las envolventes y regenerarse el expediente. No se "
            "reduce ENVCORT ni Tu para forzar cumplimiento."
        )
        for group in design.GROUP_ORDER:
            group_frames = [x["frame"] for x in failures if x["group"] == group]
            if group_frames:
                doc.add_paragraph(f"• {group}: frames {', '.join(group_frames)}")
    doc.add_paragraph(
        "Los planos son detalles académicos de coordinación y todos se rotulan NO EMITIR PARA CONSTRUCCIÓN. "
        "Incluyen tabla frame–nudos–longitud–nivel y esquema por centroides disponibles, pero el export no contiene "
        "Joint Coordinates ni una planta estructural completa; por tanto, no sustituyen planos de localización o fabricación."
    )

    add_heading(doc, "Archivos de entrega", 2)
    for text in [
        "DISENO-VIGAS-ENTREGA-FINAL-DMO.xlsx: diseño individual, grupos, fórmulas y auditoría.",
        "PLANOS-DESPIECE-VIGAS-POR-GRUPO.pdf: doce láminas, una por grupo.",
        "Planos-Despiece-Vigas/*.dxf: archivos editables para AutoCAD.",
    ]:
        doc.add_paragraph("• " + text)


def build_justification(design, records, results, groups, excel_rows):
    doc = Document()
    configure(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("JUSTIFICACIÓN DEL DISEÑO DE VIGAS\nSISTEMA DMO — R = 5")
    run.bold = True
    run.font.size = Pt(18)
    subtitle = doc.add_paragraph("Proyecto edificio residencial — Santa Marta\nAnexo técnico para memoria y sustentación")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    append_viga_chapter(doc, design, records, results, groups, excel_rows)
    doc.save(OUTPUT)


def reconcile_group_use(doc):
    """Corrige la contradicción heredada y fija residencial = Grupo I, I=1.0."""
    replacements = {
        163: ("Grupo de Uso: por tratarse de un edificio residencial de ocupación normal, el proyecto se "
              "clasifica en el Grupo de Uso I según NSR-10 A.2.5.1, con coeficiente de importancia I = 1.0."),
        530: ("Grupo I – Estructuras de ocupación normal: incluye edificaciones residenciales y demás usos "
              "que no pertenecen a los Grupos II, III o IV."),
        532: ("Grupo II – Estructuras de ocupación especial: reúne usos con condiciones especiales de "
              "ocupación o concentración según NSR-10 A.2.5.1.3."),
        538: ("Para el presente proyecto residencial, la estructura se clasifica en el Grupo de Uso I "
              "(ocupación normal)."),
        540: ("Para el Grupo de Uso I adoptado en este proyecto, el coeficiente de importancia es:"),
        541: "I = 1.0",
    }
    for index, text in replacements.items():
        if index >= len(doc.paragraphs):
            raise ValueError(f"No existe el párrafo {index} requerido para reconciliar el grupo de uso")
        doc.paragraphs[index].text = text


def _find_paragraph(doc, exact_text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"No se encontró el marcador de memoria: {exact_text}")


def _remove_body_range(doc, start_text, end_text):
    """Elimina un bloque heredado, incluyendo el encabezado inicial y sus tablas."""
    start = _find_paragraph(doc, start_text)._p
    end = _find_paragraph(doc, end_text)._p
    body = doc.element.body
    deleting = False
    for element in list(body.iterchildren()):
        if element is start:
            deleting = True
        if element is end:
            break
        if deleting:
            body.remove(element)


def _remove_stale_detail_tables(doc):
    stale_headers = {
        "Elemento|Barra|Luz libre (m)|ldh (mm)|Longitud de barra (m)",
        "Elemento|Altura efectiva d (mm)|Zona de confinamiento 2h (mm)|Separacion confinamiento, d/4 (mm)|Separacion central, d/2 (mm)|Longitud del estribo (mm)",
    }
    for table in list(doc.tables):
        if not table.rows:
            continue
        header = "|".join(cell.text.strip() for cell in table.rows[0].cells)
        if header in stale_headers:
            table._element.getparent().remove(table._element)
    stale_paragraph_starts = (
        "La longitud de corte de cada barra se obtuvo",
        "Tabla 64. Longitud de corte de las barras longitudinales",
        "La separacion de los estribos se definio",
        "Tabla 65. Separacion de estribos y zonas de confinamiento",
    )
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip().startswith(stale_paragraph_starts):
            paragraph._element.getparent().remove(paragraph._element)


def remove_legacy_beam_content(doc):
    """Retira las tablas preliminares de VC/VR para evitar contradicciones."""
    _remove_body_range(
        doc,
        "Diseño de las vigas de carga",
        "Diseño de las riostras de cimentación",
    )
    _remove_stale_detail_tables(doc)
    chapter = _find_paragraph(doc, "DISEÑO DE VIGUETAS, VIGAS Y RIOSTRAS")
    chapter.text = "DISEÑO DE VIGUETAS Y RIOSTRAS DE CIMENTACIÓN"
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("El despiece de cada tipo de vigueta, viga y riostra"):
            paragraph.text = (
                "El despiece de viguetas y riostras de cimentación se presenta en los planos "
                "correspondientes. El diseño final de las 592 vigas del modelo se incorpora en el "
                "capítulo independiente que antecede a las referencias y se sincroniza con el Excel final."
            )
            break


def _insert_appended_content_before(doc, anchor_paragraph, marker_paragraph):
    """Mueve al punto indicado el contenido agregado al final del documento."""
    body = doc.element.body
    moving = False
    elements = []
    for element in list(body.iterchildren()):
        if element is marker_paragraph._p:
            moving = True
            continue
        if moving and element.tag == qn("w:sectPr"):
            break
        if moving:
            elements.append(element)
    body.remove(marker_paragraph._p)
    for element in elements:
        body.remove(element)
        anchor_paragraph._p.addprevious(element)


def build_memory(design, records, results, groups, excel_rows):
    doc = Document(SOURCE_MEMORY) if SOURCE_MEMORY.exists() else Document()
    configure(doc)
    if SOURCE_MEMORY.exists():
        reconcile_group_use(doc)
        remove_legacy_beam_content(doc)
    reference_anchor = _find_paragraph(doc, "Referencias") if SOURCE_MEMORY.exists() else None
    marker = doc.add_paragraph("__INSERTAR_CAPITULO_VIGAS__")
    append_viga_chapter(doc, design, records, results, groups, excel_rows)
    if reference_anchor is not None:
        _insert_appended_content_before(doc, reference_anchor, marker)
    else:
        marker._element.getparent().remove(marker._element)
    doc.save(MEMORY_OUTPUT)


def main():
    design = load_design_module()
    records, _ = design.extract_inputs()
    results = [design.design_record(record) for record in records]
    groups = design.build_group_records(records, results)
    excel_rows = read_excel_group_rows()
    validate_excel_rows(design, groups, excel_rows)
    build_justification(design, records, results, groups, excel_rows)
    build_memory(design, records, results, groups, excel_rows)
    print(f"OK: {OUTPUT}")
    print(f"OK: {MEMORY_OUTPUT}")
    print(f"Vigas: {len(results)}; cumplen: {sum(x['overall'] for x in results)}; no cumplen: {sum(not x['overall'] for x in results)}")


if __name__ == "__main__":
    main()
