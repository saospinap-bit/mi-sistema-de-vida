#!/usr/bin/env python3
"""Genera una guía de sustentación del Excel reconciliado de diseño de vigas.

La guía se construye desde el libro vigente para que cantidades, demandas, estados,
ejemplos y tablas no se transcriban manualmente. Las ecuaciones se insertan como
objetos OMML nativos, editables con el editor de ecuaciones de Microsoft Word.
"""
from __future__ import annotations

from collections import Counter, defaultdict
from importlib.metadata import PackageNotFoundError, version
from pathlib import Path
import platform
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

BASE = Path(__file__).resolve().parent
DESIGN = BASE / "Diseño de vigas proyecto diseño DEF.xlsx"
SAP = BASE / "resultados sap.xlsx"
GEOMETRY = BASE / "geomatria sap.xlsx"
OUTPUT = BASE / "GUIA_SUSTENTACION_EXCEL_DISENO_DE_VIGAS.docx"
PREVIEW = BASE / "Planos-Autocad-B2-Vigas-Continuas" / "VISTA-PREVIA-TODAS-LAS-LAMINAS.png"

DETAIL = "Todas las Vigas (592)"
SUMMARIES = ("Vigas de Carga (7)", "Vigas de Rigidez (5)")
GROUP_ORDER = ["VC1", "VC2", "VC3", "VC4", "VC5", "VC6", "VC7", "VR1", "VR N1", "VR2", "VR3", "VRAUX"]
EXPECTED_FRAMES = 592
EXPECTED_GROUPS = 12
EXPECTED_FORMULAS = 21_744

NAVY = "17365D"
BLUE = "2F75B5"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EAF2F8"
GREEN = "C6E0B4"
PALE_GREEN = "E2F0D9"
YELLOW = "FFF2CC"
RED = "F4CCCC"
DARK_RED = "9C0006"
GRAY = "E7E6E6"
DARK_GRAY = "595959"
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)


def clean(value) -> str:
    return " ".join(str(value).split())


def fmt(value, decimals=1) -> str:
    return f"{float(value):.{decimals}f}"


def package_version(distribution: str) -> str:
    try:
        return version(distribution)
    except PackageNotFoundError:
        return "no disponible"


def _set_cell_shading(cell, fill: str):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=55, start=60, bottom=55, end=60):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _repeat_header(row):
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _set_repeat_table_rows(row, cant_split=True):
    properties = row._tr.get_or_add_trPr()
    if cant_split:
        node = OxmlElement("w:cantSplit")
        properties.append(node)


def _set_cell_text(cell, text, *, bold=False, color=BLACK, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = str(text)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_margins(cell)
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(size)
            run.bold = bold
            run.font.color.rgb = color


def add_table(document, headers, rows, widths=None, status_columns=(), font_size=8.0):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_text(cell, header, bold=True, color=WHITE, size=font_size)
        _set_cell_shading(cell, NAVY)
    _repeat_header(table.rows[0])
    for values in rows:
        row = table.add_row()
        _set_repeat_table_rows(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            _set_cell_text(cell, value, size=font_size)
            if index in status_columns:
                text = str(value)
                _set_cell_shading(cell, RED if "NO CUMPLE" in text else PALE_GREEN)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    if "NO CUMPLE" in text:
                        run.font.color.rgb = RGBColor(156, 0, 6)
        if widths:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    if widths:
        for index, width in enumerate(widths):
            table.rows[0].cells[index].width = Cm(width)
    document.add_paragraph()
    return table


def set_repeat_update_fields(document):
    settings = document.settings._element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend((field_begin, instruction, field_end))


def add_toc(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Actualice la tabla de contenido en Word con clic derecho > Actualizar campo."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, placeholder, end))


def configure_document(document):
    set_repeat_update_fields(document)
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)
    header = section.header.paragraphs[0]
    header.text = "GUÍA DE SUSTENTACIÓN — EXCEL DE DISEÑO DE VIGAS"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(89, 89, 89)
    add_page_number(section.footer.paragraphs[0])

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in (("Title", 25, NAVY), ("Heading 1", 17, NAVY), ("Heading 2", 13, BLUE), ("Heading 3", 11, DARK_GRAY)):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(9)
        style.paragraph_format.space_after = Pt(5)


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(document, text, *, bold_prefix=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)
    return paragraph


def add_bullet(document, text, level=0):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    paragraph.paragraph_format.first_line_indent = Cm(-0.35)
    paragraph.add_run("• ").bold = True
    paragraph.add_run(text)
    return paragraph


def add_numbered(document, number, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.7)
    paragraph.paragraph_format.first_line_indent = Cm(-0.7)
    run = paragraph.add_run(f"{number}. ")
    run.bold = True
    paragraph.add_run(text)
    return paragraph


def add_callout(document, title, text, fill=YELLOW):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_shading(cell, fill)
    _set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(f"{title}: ")
    run.bold = True
    paragraph.add_run(text)
    document.add_paragraph()
    return table


def add_flow(document, steps):
    table = document.add_table(rows=1, cols=len(steps))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, (title, detail) in enumerate(steps):
        cell = table.cell(0, index)
        _set_cell_shading(cell, LIGHT_BLUE if index % 2 == 0 else PALE_GREEN)
        _set_cell_margins(cell, top=100, start=70, bottom=100, end=70)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY)
        p.add_run(f"\n{detail}")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    document.add_paragraph()
    return table


# ---------- Objetos matemáticos nativos de Word (OMML) ----------
def math_run(text):
    run = OxmlElement("m:r")
    token = OxmlElement("m:t")
    token.text = str(text)
    run.append(token)
    return run


def append_math(parent, parts):
    if isinstance(parts, (str, int, float)):
        parent.append(math_run(parts))
    elif getattr(parts, "tag", None) is not None:
        parent.append(parts)
    else:
        for part in parts:
            append_math(parent, part)


def math_box(tag, parts):
    box = OxmlElement(tag)
    append_math(box, parts)
    return box


def frac(numerator, denominator):
    node = OxmlElement("m:f")
    node.append(math_box("m:num", numerator))
    node.append(math_box("m:den", denominator))
    return node


def sub(base, index):
    node = OxmlElement("m:sSub")
    node.append(math_box("m:e", base))
    node.append(math_box("m:sub", index))
    return node


def sup(base, power):
    node = OxmlElement("m:sSup")
    node.append(math_box("m:e", base))
    node.append(math_box("m:sup", power))
    return node


def radical(parts):
    node = OxmlElement("m:rad")
    properties = OxmlElement("m:radPr")
    hide = OxmlElement("m:degHide")
    hide.set(qn("m:val"), "1")
    properties.append(hide)
    node.append(properties)
    node.append(OxmlElement("m:deg"))
    node.append(math_box("m:e", parts))
    return node


def cases(lines):
    node = OxmlElement("m:d")
    properties = OxmlElement("m:dPr")
    begin = OxmlElement("m:begChr")
    begin.set(qn("m:val"), "{")
    end = OxmlElement("m:endChr")
    end.set(qn("m:val"), "")
    properties.extend((begin, end))
    node.append(properties)
    matrix = OxmlElement("m:m")
    for line in lines:
        row = OxmlElement("m:mr")
        cell = OxmlElement("m:e")
        append_math(cell, line)
        row.append(cell)
        matrix.append(row)
    expression = OxmlElement("m:e")
    expression.append(matrix)
    node.append(expression)
    return node


def add_equation(document, parts, label=None):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(5)
    math_paragraph = OxmlElement("m:oMathPara")
    properties = OxmlElement("m:oMathParaPr")
    justification = OxmlElement("m:jc")
    justification.set(qn("m:val"), "centerGroup")
    properties.append(justification)
    math_paragraph.append(properties)
    equation = OxmlElement("m:oMath")
    append_math(equation, parts)
    math_paragraph.append(equation)
    paragraph._p.append(math_paragraph)
    if label:
        caption = document.add_paragraph(label)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.italic = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(89, 89, 89)
    return paragraph


def portrait_section(document):
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    return section


def landscape_section(document):
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    return section


# ---------- Extracción y validación ----------
def sap_rows(workbook, sheet_name):
    sheet = workbook[sheet_name]
    headers = [cell.value for cell in sheet[2]]
    for values in sheet.iter_rows(min_row=4, values_only=True):
        if values and values[0] is not None:
            yield dict(zip(headers, values))


def load_data():
    with ZipFile(DESIGN) as archive:
        names = archive.namelist()
        if any("externalLinks" in name for name in names):
            raise ValueError("El libro conserva externalLinks")
        worksheet_xml = [
            archive.read(name) for name in names
            if name.startswith("xl/worksheets/sheet") and name.endswith(".xml")
        ]
    formula_count = sum(xml.count(b"<f") for xml in worksheet_xml)
    payload = b"".join(worksheet_xml)
    if formula_count != EXPECTED_FORMULAS or b"[1]" in payload or b't="e"' in payload:
        raise ValueError("El libro no supera la validación OOXML")

    formulas = load_workbook(DESIGN, read_only=True, data_only=False, keep_links=False)
    values = load_workbook(DESIGN, read_only=True, data_only=True, keep_links=False)
    parameters = [
        (values["Parametros"].cell(row, 1).value, values["Parametros"].cell(row, 2).value, values["Parametros"].cell(row, 3).value)
        for row in range(2, 13)
    ]
    detail_f = formulas[DETAIL]
    detail_v = values[DETAIL]
    headers = [cell.value for cell in detail_v[3]]
    index = {header: i for i, header in enumerate(headers)}

    rows = []
    frames = set()
    statuses = defaultdict(Counter)
    failures = defaultdict(Counter)
    checks = ["Flexión −", "Flexión +", "Sección cortante", "Cortante", "Chequeo torsión", "Chequeo V-T", "Chequeo Al"]
    for values_row in detail_v.iter_rows(min_row=4, values_only=True):
        if values_row[0] is None:
            continue
        row = dict(zip(headers, values_row))
        frame = str(row["Frame"])
        if frame in frames:
            raise ValueError(f"Frame repetido: {frame}")
        frames.add(frame)
        rows.append(row)
        statuses[str(row["Grupo"])][str(row["ESTADO"])] += 1
        for check in checks:
            if row[check] not in ("CUMPLE", "DESPRECIABLE"):
                failures[str(row["Grupo"])][check] += 1

    summaries = {}
    summary_formulas = {}
    for sheet_name in SUMMARIES:
        sheet_v = values[sheet_name]
        sheet_f = formulas[sheet_name]
        sheet_headers = [cell.value for cell in sheet_v[3]]
        for row_number, values_row in enumerate(sheet_v.iter_rows(min_row=4, values_only=True), 4):
            if values_row[0] is None:
                continue
            row = dict(zip(sheet_headers, values_row))
            group = str(row["Grupo"])
            summaries[group] = row
            summary_formulas[group] = sheet_f.cell(row_number, 52).value

    formula_by_sheet = {}
    for sheet in formulas.worksheets:
        formula_by_sheet[sheet.title] = sum(
            1 for row in sheet.iter_rows() for cell in row if cell.data_type == "f"
        )
    detail_formula_example = {
        get_column_letter(column): detail_f.cell(4, column).value
        for column in range(1, 63) if detail_f.cell(4, column).data_type == "f"
    }
    detail_status_formula = detail_f["AZ4"].value
    formulas.close()
    values.close()

    sap = load_workbook(SAP, read_only=True, data_only=True)
    aliases = {"VR1 N1": "VR N1", "VR AUX": "VRAUX"}
    assignments = defaultdict(set)
    owners = {}
    for row in sap_rows(sap, "Groups 2 - Assignments"):
        if str(row.get("ObjectType")) != "Frame":
            continue
        raw = str(row.get("GroupName")).strip()
        group = aliases.get(raw, raw)
        if group not in GROUP_ORDER:
            continue
        frame = str(row["ObjectLabel"])
        if frame in owners and owners[frame] != group:
            raise ValueError(f"Frame {frame} asignado a dos grupos")
        owners[frame] = group
        assignments[group].add(frame)
    force_rows = Counter()
    for row in sap_rows(sap, "Element Forces - Frames"):
        frame = str(row["Frame"])
        case = str(row["OutputCase"])
        if frame in owners and case in ("ENVFLEX", "ENVCORT"):
            force_rows[case] += 1
    sap.close()

    if len(frames) != EXPECTED_FRAMES or len(owners) != EXPECTED_FRAMES or set(summaries) != set(GROUP_ORDER):
        raise ValueError("La fuente y el libro no contienen los 592 frames/12 grupos esperados")
    if frames != set(owners):
        missing_in_sap = sorted(frames - set(owners))
        missing_in_excel = sorted(set(owners) - frames)
        raise ValueError(
            "No coinciden las identidades de frame SAP–Excel: "
            f"solo Excel={missing_in_sap[:10]}, solo SAP={missing_in_excel[:10]}"
        )
    for row in rows:
        frame = str(row["Frame"])
        excel_group = str(row["Grupo"])
        if owners[frame] != excel_group:
            raise ValueError(
                f"Grupo inconsistente para frame {frame}: SAP={owners[frame]}, Excel={excel_group}"
            )
    for group in GROUP_ORDER:
        if len(assignments[group]) != sum(statuses[group].values()):
            raise ValueError(f"Cantidad inconsistente en {group}")

    detail_totals = Counter()
    for counter in statuses.values():
        detail_totals.update(counter)
    summary_totals = Counter(str(row["ESTADO"]) for row in summaries.values())
    summary_compliant_groups = [
        group for group in GROUP_ORDER if str(summaries[group]["ESTADO"]) == "CUMPLE"
    ]
    summary_noncompliant_groups = [
        group for group in GROUP_ORDER if str(summaries[group]["ESTADO"]) == "NO CUMPLE"
    ]

    return {
        "parameters": parameters,
        "headers": headers,
        "rows": rows,
        "summaries": summaries,
        "statuses": statuses,
        "detail_totals": detail_totals,
        "summary_totals": summary_totals,
        "summary_compliant_groups": summary_compliant_groups,
        "summary_noncompliant_groups": summary_noncompliant_groups,
        "failures": failures,
        "assignments": assignments,
        "force_rows": force_rows,
        "formula_count": formula_count,
        "formula_by_sheet": formula_by_sheet,
        "environment": {
            "Python": platform.python_version(),
            "python-docx": package_version("python-docx"),
            "openpyxl": package_version("openpyxl"),
            "Apache POI": "5.4.1 (proceso de reconciliación)",
        },
        "detail_formula_example": detail_formula_example,
        "detail_status_formula": detail_status_formula,
        "summary_status_formula": next(iter(summary_formulas.values())),
    }


def group_failure_labels(row):
    mapping = [
        ("Flexión −", "flexión negativa"),
        ("Flexión +", "flexión positiva"),
        ("Sección cortante", "límite de sección a cortante"),
        ("Cortante", "resistencia a cortante"),
        ("Chequeo torsión", "refuerzo de torsión"),
        ("Chequeo V-T", "interacción V-T"),
        ("Chequeo Al", "acero longitudinal de torsión"),
    ]
    labels = [label for field, label in mapping if row[field] not in ("CUMPLE", "DESPRECIABLE")]
    return ", ".join(labels) if labels else "—"


def find_frame(rows, frame):
    return next(row for row in rows if str(row["Frame"]) == str(frame))


def add_cover(document):
    document.add_paragraph().paragraph_format.space_after = Pt(28)
    title = document.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("GUÍA DE SUSTENTACIÓN\nDEL EXCEL DE DISEÑO DE VIGAS")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Edificio residencial — Santa Marta\n592 frames · 12 grupos · cálculo por flexión, cortante, capacidad y torsión")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    document.add_paragraph().paragraph_format.space_after = Pt(18)
    add_callout(
        document,
        "Propósito",
        "servir como libreto técnico para explicar mañana qué contiene el Excel, de dónde salen los datos, cómo se agruparon las vigas, qué calcula cada bloque, cómo se interpretan los estados y cuáles son los límites de la entrega.",
        LIGHT_BLUE,
    )
    table = document.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cover_rows = [
        ("Proyecto", "Edificio residencial Santa Marta — Grupo 6"),
        ("Libro explicado", "Diseño de vigas proyecto diseño DEF.xlsx"),
        ("Fuentes", "resultados sap.xlsx + geomatria sap.xlsx"),
        ("Fecha de preparación", "18 de julio de 2026"),
        ("Alcance", "Entrega académica; no emitir para construcción"),
    ]
    for index, (label, value) in enumerate(cover_rows):
        _set_cell_text(table.cell(index, 0), label, bold=True, color=WHITE, size=9)
        _set_cell_shading(table.cell(index, 0), NAVY)
        _set_cell_text(table.cell(index, 1), value, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
    document.add_paragraph()
    warning = document.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = warning.add_run("DETALLE ACADÉMICO — NO EMITIR PARA CONSTRUCCIÓN")
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    run.font.size = Pt(12)
    document.add_page_break()


def add_executive_script(document, data):
    detail = data["detail_totals"]
    summary = data["summary_totals"]
    add_heading(document, "1. Guion ejecutivo para una exposición de 12 a 15 minutos", 1)
    add_callout(
        document,
        "Idea central",
        "El Excel no inventa las demandas. Reorganiza resultados vigentes de SAP2000, conserva la trazabilidad por frame, calcula las verificaciones y separa claramente el diseño individual de 592 frames del diseño común resumido en 12 grupos.",
        PALE_GREEN,
    )
    script = [
        ("0:00–1:00", "Objetivo", "Presentar el flujo SAP → reconciliación → Excel → comprobaciones → CAD, aclarando que es una entrega académica."),
        ("1:00–3:00", "Fuentes", "Mostrar que las fuerzas vienen de ENVFLEX/ENVCORT, la geometría de geomatria sap.xlsx y las familias de Groups 2 - Assignments."),
        ("3:00–5:00", "Agrupación", "Explicar los 240 frames VC y 352 frames VR/VRAUX, sin afirmar que todas las vigas de un grupo tienen la misma luz."),
        ("5:00–7:00", "Arquitectura", "Recorrer Parametros, dos resúmenes y Todas las Vigas (592); explicar A:Q, decisiones de armado y 36 fórmulas por fila."),
        ("7:00–10:00", "Cálculo", "Explicar flexión, cortante por capacidad, torsión e interacción V-T con las ecuaciones nativas de esta guía."),
        ("10:00–12:00", "Ejemplos", "Mostrar VC1 como caso CUMPLE y VC5/frame 84 como caso NO CUMPLE."),
        ("12:00–14:00", "Resultados", f"Diferenciar {detail['CUMPLE']}/{detail['NO CUMPLE']} estados detallados de {summary['CUMPLE']}/{summary['NO CUMPLE']} estados resumen y explicar por qué no son contradictorios."),
        ("14:00–15:00", "Cierre", "Resaltar trazabilidad, resultados almacenados después del recálculo, cero vínculos externos y necesidad de revisión profesional de los incumplimientos."),
    ]
    add_table(document, ("Tiempo", "Tema", "Qué decir"), script, widths=(2.2, 3.0, 11.5), font_size=8.5)
    add_heading(document, "Frase de apertura sugerida", 2)
    add_body(
        document,
        "“El objetivo del libro fue convertir los resultados del modelo SAP2000 en un diseño de vigas trazable. Para ello se conservaron los 592 identificadores de frame, se respetaron los 12 grupos asignados en SAP, se reconstruyeron las demandas desde las envolventes vigentes y se ejecutaron 21.744 fórmulas de flexión, cortante, capacidad y torsión. El resultado no se presenta como una aprobación: se conservan y se explican todos los NO CUMPLE.”",
    )
    add_heading(document, "Frase de cierre sugerida", 2)
    add_body(
        document,
        f"“La principal fortaleza del Excel es la trazabilidad: cada demanda se puede regresar al frame SAP que la originó. La principal conclusión técnica es que el cálculo individual y el resumen común por grupo deben leerse por separado; existen {detail['NO CUMPLE']} frames detallados NO CUMPLE y {summary['NO CUMPLE']} armados resumen que requieren rediseño antes de cualquier revisión profesional o emisión constructiva.”",
    )


def add_workflow_and_sources(document, data):
    add_heading(document, "2. Qué se hizo y de dónde salió la información", 1)
    add_flow(document, [
        ("1. SAP2000", "Geometría, grupos y fuerzas"),
        ("2. Python", "Reconciliación y controles"),
        ("3. Excel", "21.744 fórmulas y valores almacenados"),
        ("4. CAD", "Demandas y estados por frame"),
    ])
    add_heading(document, "2.1 Fuentes utilizadas", 2)
    rows = [
        ("resultados sap.xlsx", "Groups 2 - Assignments", "Asignación de cada frame a VC1–VC7, VR1, VR N1, VR2, VR3 o VRAUX."),
        ("resultados sap.xlsx", "Element Forces - Frames", "Estaciones y envolventes Max/Min de M3, V2 y T para ENVFLEX y ENVCORT."),
        ("resultados sap.xlsx", "Frame Props 01 - General", "Dimensiones t2 y t3 de las secciones del modelo."),
        ("geomatria sap.xlsx", "Connectivity - Frame", "Longitud y conectividad I–J de cada frame."),
        ("geomatria sap.xlsx", "Frame Section Assignments", "Sección analítica asociada con cada frame."),
        ("Diseño de vigas…xlsx", "Armados adoptados", "Número de barras, capas, estribos, ramas y barras dedicadas a torsión que luego se verifican."),
    ]
    add_table(document, ("Archivo", "Hoja", "Uso"), rows, widths=(4.0, 4.0, 8.7), font_size=8.4)
    add_body(
        document,
        f"Para las vigas se procesaron {data['force_rows']['ENVFLEX']:,} registros de ENVFLEX y "
        f"{data['force_rows']['ENVCORT']:,} registros de ENVCORT. Cada registro corresponde a una estación del frame y a un extremo Max o Min de la envolvente.".replace(",", "."),
    )
    add_callout(
        document,
        "Punto que debes defender",
        "Las fuerzas no fueron digitadas una por una ni tomadas de capturas. Se agregaron directamente desde la tabla Element Forces - Frames, evitando perder estaciones críticas o confundir frames.",
        YELLOW,
    )
    add_heading(document, "2.2 Componentes de fuerza utilizados", 2)
    add_bullet(document, "M3: momento flector principal del elemento de viga según sus ejes locales.")
    add_bullet(document, "V2: cortante principal asociado con la flexión M3.")
    add_bullet(document, "T: torsión alrededor del eje longitudinal del frame.")
    add_bullet(document, "ENVFLEX: envolvente usada para Mu−, Mu+ y Vu de flexión.")
    add_bullet(document, "ENVCORT: envolvente conservadora usada para Vu de diseño y Tu.")
    add_callout(
        document,
        "Si preguntan por los ejes locales",
        "debes responder que M3/V2/T son los componentes exportados para los elementos de viga y que su orientación debe verificarse con los ejes locales del modelo; el script no cambia signos arbitrariamente, sino que separa el máximo positivo y la magnitud del mínimo negativo.",
        LIGHT_BLUE,
    )


def add_grouping(document, data):
    add_heading(document, "3. Cómo se agruparon las vigas", 1)
    add_body(
        document,
        "La agrupación no fue inferida por el script a partir de una semejanza geométrica. Se leyó directamente de la hoja Groups 2 - Assignments de SAP2000. El script normalizó únicamente dos nombres equivalentes —VR1 N1 a VR N1 y VR AUX a VRAUX— y luego exigió que cada frame perteneciera a un solo grupo de vigas.",
    )
    counts = [(group, len(data["assignments"][group]), "Carga" if group.startswith("VC") else "Rigidez/auxiliar") for group in GROUP_ORDER]
    counts.append(("TOTAL", sum(row[1] for row in counts), "12 grupos"))
    add_table(document, ("Grupo SAP", "Frames", "Familia"), counts, widths=(4.0, 3.0, 7.0), font_size=9)
    add_body(document, "Las siete familias VC contienen 240 frames y las cinco familias VR/VRAUX contienen 352 frames; en total se verificaron 592 frames sin duplicados.")
    add_heading(document, "3.1 Qué representa un grupo", 2)
    add_bullet(document, "Una familia de diseño definida en SAP, no necesariamente una única longitud.")
    add_bullet(document, "Una sección adoptada común en el resumen del grupo.")
    add_bullet(document, "Una envolvente grupal obtenida tomando máximos entre todos sus frames.")
    add_bullet(document, "Un armado tipo resumido que se evalúa contra esa envolvente conservadora.")
    add_heading(document, "3.2 Qué no debes afirmar", 2)
    add_callout(
        document,
        "No digas",
        "“todas las vigas de cada grupo tienen la misma longitud y los mismos apoyos”. Hay grupos con rangos amplios —por ejemplo VC7 y VR N1—. Lo correcto es decir que son grupos SAP y que las columnas D y E conservan la longitud mínima y máxima de cada familia.",
        RED,
    )
    add_heading(document, "3.3 Cómo se construyó la envolvente de cada frame", 2)
    add_equation(document, (sub("M", "u"), "−", "=", "max", "(", "0", ",", "−", "min", "(", sub("M", "3,ENVFLEX"), ")", ")"))
    add_equation(document, (sub("M", "u"), "+", "=", "max", "(", "0", ",", "max", "(", sub("M", "3,ENVFLEX"), ")", ")"))
    add_equation(document, (sub("V", "u,F"), "=", "max", "|", sub("V", "2,ENVFLEX"), "|"))
    add_equation(document, (sub("V", "u,C"), "=", "max", "|", sub("V", "2,ENVCORT"), "|"))
    add_equation(document, (sub("T", "u"), "=", "max", "|", sub("T", "ENVCORT"), "|"))
    add_body(document, "En el resumen grupal se guarda además el frame que controla Mu−, Mu+ y Vu. La columna B del resumen conserva el frame controlador de la torsión, mientras K, M y P conservan los controladores de momento negativo, positivo y cortante.")


def add_architecture(document, data):
    add_heading(document, "4. Arquitectura del Excel", 1)
    rows = [
        ("Parametros", "11 parámetros", "Materiales, factores φ, recubrimiento, área y diámetro de barra #5. Los parámetros están internalizados."),
        ("Vigas de Carga (7)", "7 filas de diseño", "Envolventes y armado común de VC1 a VC7."),
        ("Vigas de Rigidez (5)", "5 filas de diseño", "Envolventes y armado común de VR1, VR N1, VR2, VR3 y VRAUX."),
        ("Todas las Vigas (592)", "592 filas", "Una fila por frame, con demanda, armado individual y estado."),
    ]
    add_table(document, ("Hoja", "Contenido", "Función"), rows, widths=(4.2, 3.2, 9.3), font_size=8.5)
    add_heading(document, "4.1 Tres tipos de celdas en cada fila de diseño", 2)
    add_body(document, "Cada fila de las tres hojas de diseño tiene 62 columnas (A:BJ), distribuidas de la siguiente forma:")
    add_table(
        document,
        ("Tipo", "Cantidad", "Columnas", "Ejemplos"),
        (
            ("Datos reconciliados", "17", "A:Q", "Grupo, frame, luces, sección, Mu, Vu y Tu."),
            ("Decisiones de armado", "9", "S, T, Z, AA, AN, AO, AX, BA, BB", "Barras, capas, estribo, ramas, barras torsionales, db y área de rama."),
            ("Fórmulas", "36", "R:BJ, excepto decisiones", "As, d, φMn, Mpr, Vcap, φVn, torsión, interacción y estado."),
        ),
        widths=(3.6, 2.0, 4.5, 6.6),
        font_size=8.5,
    )
    add_callout(
        document,
        "Respuesta clave",
        "El Excel no selecciona automáticamente todo el armado. Hay nueve celdas de decisión por fila. El libro calcula si ese armado adoptado satisface la demanda. El script de reconciliación actualiza A:Q y las referencias; no aumenta barras para fabricar un CUMPLE.",
        YELLOW,
    )
    add_heading(document, "4.2 Parámetros internos", 2)
    parameter_rows = []
    for label, value, unit in data["parameters"]:
        if label == "Reservado":
            continue
        parameter_rows.append((label, "—" if value is None else value, unit))
    add_table(document, ("Parámetro", "Valor", "Unidad/nota"), parameter_rows, widths=(7.0, 3.0, 6.0), font_size=8.8)
    add_body(document, "El diámetro de estribo B9 = 9.5 mm es solo una referencia. Las fórmulas usan BA de la misma fila, porque existen estribos #3 de 9.5 mm y #4 de 12.7 mm.")


def add_column_map(document, data):
    landscape_section(document)
    add_heading(document, "5. Mapa completo de columnas A:BJ", 1)
    groups = [
        ("Identificación y demandas", 1, 17, "Datos SAP/geométricos reconciliados"),
        ("Flexión negativa", 18, 24, "Acero superior y φMn−"),
        ("Flexión positiva", 25, 31, "Acero inferior y φMn+"),
        ("Capacidad y cortante", 32, 45, "Mpr, Vcap, Vu diseño, Vc, Vs, estribos y φVn"),
        ("Torsión y estado", 46, 52, "Tth, At/s, Al, barras torsionales y estado"),
        ("Interacción y cierre", 53, 62, "db, Av/s, interacción V-T y chequeo Al"),
    ]
    rows = []
    for block, first, last, purpose in groups:
        names = [f"{get_column_letter(column)} — {data['headers'][column-1]}" for column in range(first, last + 1)]
        rows.append((block, "\n".join(names), purpose))
    add_table(document, ("Bloque", "Columnas", "Qué demuestra"), rows, widths=(5.0, 13.0, 7.5), font_size=7.4)
    add_heading(document, "5.1 Cómo leer una fila de izquierda a derecha", 2)
    for number, text in enumerate((
        "Identifique grupo y frame en A:B.",
        "Revise geometría y sección en D:I.",
        "Ubique las demandas en J:Q y sus frames controladores cuando esté en una hoja resumen.",
        "Compruebe el acero superior en R:X y el inferior en Y:AE.",
        "Siga Mpr, Vcap y cortante en AF:AS.",
        "Revise torsión y estado en AT:AZ.",
        "Cierre con interacción V-T y acero torsional en BA:BJ.",
    ), 1):
        add_numbered(document, number, text)
    portrait_section(document)


def add_formulas(document):
    add_heading(document, "6. Sustento de las fórmulas", 1)
    add_callout(document, "Unidades", "f'c, fy y fyt en MPa (N/mm²); b, h, d, db y áreas en mm/mm²; fuerzas en kN; momentos y torsión en kN·m. Por eso aparecen conversiones de 10³ y 10⁶.", LIGHT_BLUE)

    add_heading(document, "6.1 Flexión negativa y positiva", 2)
    add_body(document, "El procedimiento es el mismo para ambas caras; cambia la demanda y el acero superior o inferior.")
    add_equation(document, ("d", "=", "h", "−", sub("c", "c"), "−", sub("d", "b,est"), "−", frac(sub("d", "b,long"), "2")), "Peralte efectivo para una capa")
    add_equation(document, (sub("R", "n"), "=", frac((sub("M", "u"), "·", sup("10", "6")), (sub("φ", "f"), "·", "b", "·", sup("d", "2")))), "Parámetro resistente")
    add_equation(document, ("ρ", "=", frac(("0.85", "·", "f′c"), sub("f", "y")), "·", "[", "1", "−", radical(("1", "−", frac(("2", "·", sub("R", "n")), ("0.85", "·", "f′c")))), "]"), "Cuantía requerida")
    add_equation(document, (sub("A", "s,min"), "=", "max", "[", frac("1.4", sub("f", "y")), ",", frac(("0.25", "·", radical("f′c")), sub("f", "y")), "]", "·", "b", "·", "d"), "Acero mínimo")
    add_equation(document, (sub("A", "s,req"), "=", "max", "(", "ρ", "·", "b", "·", "d", ",", sub("A", "s,min"), ")"), "Acero requerido")
    add_equation(document, (sub("A", "s,prov"), "=", sub("n", "b"), "·", sub("A", "b")), "Acero proporcionado")
    add_equation(document, ("a", "=", frac((sub("A", "s,prov"), "·", sub("f", "y")), ("0.85", "·", "f′c", "·", "b"))), "Bloque equivalente de compresión")
    add_equation(document, (sub("φ", "f"), sub("M", "n"), "=", frac((sub("φ", "f"), "·", sub("A", "s,prov"), "·", sub("f", "y"), "·", "(", "d", "−", frac("a", "2"), ")"), sup("10", "6")), "≥", sub("M", "u")), "Condición de flexión")
    add_body(document, "Si hay más de una capa, U o AB calcula un centroide ponderado. La separación vertical de capas usa 25 mm y el diámetro longitudinal #5.")

    add_heading(document, "6.2 Cortante por capacidad", 2)
    add_equation(document, (sub("M", "pr"), "=", frac(("1.25", "·", sub("A", "s"), "·", sub("f", "y"), "·", "(", "d", "−", frac(sub("a", "pr"), "2"), ")"), sup("10", "6"))), "Momento probable")
    add_equation(document, (sub("V", "cap"), "=", sub("V", "u,ENVFLEX"), "+", frac((sup(sub("M", "pr"), "−"), "+", sup(sub("M", "pr"), "+")), sub("L", "n"))), "Cortante por capacidad")
    add_equation(document, (sub("V", "u,diseño"), "=", "max", "(", sub("V", "u,ENVCORT"), ",", sub("V", "cap"), ")"), "Demanda usada en cortante")
    add_equation(document, (sub("V", "c"), "=", frac(("0.17", "·", radical("f′c"), "·", "b", "·", "d"), sup("10", "3"))), "Aporte del concreto, kN")
    add_equation(document, (sub("V", "s,req"), "=", "max", "(", frac(sub("V", "u,diseño"), sub("φ", "v")), "−", sub("V", "c"), ",", "0", ")"), "Aporte requerido del acero")
    add_equation(document, (sub("V", "s,max"), "=", frac(("0.66", "·", radical("f′c"), "·", "b", "·", "d"), sup("10", "3"))), "Límite de Vs usado en AM")
    add_equation(document, (sub("φ", "v"), sub("V", "n"), "=", sub("φ", "v"), "·", "[", sub("V", "c"), "+", frac((frac(sub("A", "v"), "s"), "·", sub("f", "yt"), "·", "d"), sup("10", "3")), "]", "≥", sub("V", "u,diseño")), "Resistencia a cortante")
    add_body(document, "Las separaciones AP y AQ son el mínimo de resistencia, límites geométricos/DMO, diámetro de barra longitudinal, diámetro de estribo y requisitos de torsión; después se redondean hacia abajo a múltiplos de 10 mm.")

    add_heading(document, "6.3 Torsión", 2)
    add_equation(document, (sub("A", "cp"), "=", "b", "·", "h", "   ;   ", sub("p", "cp"), "=", "2", "·", "(", "b", "+", "h", ")"), "Geometría bruta")
    add_equation(document, (sub("φ", "t"), sub("T", "th"), "=", frac((sub("φ", "t"), "·", "0.083", "·", radical("f′c"), "·", sup(sub("A", "cp"), "2")), (sub("p", "cp"), "·", sup("10", "6")))), "Umbral de torsión")
    add_equation(document, (sub("T", "u"), "≤", sub("φ", "t"), sub("T", "th"), "  →  DESPRECIABLE"), "Decisión inicial")
    add_equation(document, (sub("A", "oh"), "=", "[", "b", "−", "2", "·", "(", sub("c", "c"), "+", frac(sub("d", "b,est"), "2"), ")", "]", "·", "[", "h", "−", "2", "·", "(", sub("c", "c"), "+", frac(sub("d", "b,est"), "2"), ")", "]"), "Área encerrada por el eje del estribo")
    add_equation(document, (sub("A", "o"), "=", "0.85", "·", sub("A", "oh")), "Área efectiva de torsión")
    add_equation(document, (frac(sub("A", "t"), "s"), "=", frac((sub("T", "u"), "·", sup("10", "6")), (sub("φ", "t"), "·", "2", "·", sub("A", "o"), "·", sub("f", "yt")))), "Refuerzo transversal torsional para θ = 45°")
    add_equation(document, (sub("A", "l"), "=", frac(sub("A", "t"), "s"), "·", sub("p", "h"), "·", frac(sub("f", "yt"), sub("f", "y"))), "Acero longitudinal de torsión")

    add_heading(document, "6.4 Interacción cortante–torsión", 2)
    add_equation(
        document,
        (
            radical((
                sup(frac((sub("V", "u"), "·", sup("10", "3")), ("b", "·", "d")), "2"), "+",
                sup(frac((sub("T", "u"), "·", sup("10", "6"), "·", sub("p", "h")), ("1.7", "·", sup(sub("A", "oh"), "2"))), "2"),
            )), "≤", sub("φ", "v"), "·", "[", frac((sub("V", "c"), "·", sup("10", "3")), ("b", "·", "d")), "+", "0.66", "·", radical("f′c"), "]",
        ),
        "Interacción V-T de BF ≤ BG",
    )
    add_heading(document, "6.5 Lógica del estado", 2)
    add_equation(document, ("ESTADO", "=", cases((("CUMPLE, si X, AE, AS, AY, BH y BJ son aceptables"), ("NO CUMPLE, en cualquier otro caso")))), "Lógica conceptual de AZ")
    add_body(document, "En detalle, AY puede ser CUMPLE o DESPRECIABLE. En los resúmenes se añade una condición: deben existir frames del grupo y ninguno puede tener AZ distinto de CUMPLE.")


def add_worked_example(document, data):
    add_heading(document, "7. Ejemplo completo para exponer: grupo VC1", 1)
    row = data["summaries"]["VC1"]
    add_body(document, "VC1 es un buen ejemplo porque permite recorrer todo el procedimiento y su estado final es CUMPLE.")
    inputs = [
        ("Sección", f"{int(row['b adopt. (mm)'])} × {int(row['h adopt. (mm)'])} mm"),
        ("Longitud", f"{fmt(row['Ln adopt. mín (m)'],2)} m"),
        ("Mu− / frame", f"{fmt(row['Mu− ENVFLEX (kN·m)'],4)} kN·m / F{row['Frame Mu−']}"),
        ("Mu+ / frame", f"{fmt(row['Mu+ ENVFLEX (kN·m)'],4)} kN·m / F{row['Frame Mu+']}"),
        ("Vu ENVFLEX / ENVCORT", f"{fmt(row['Vu ENVFLEX (kN)'],3)} / {fmt(row['Vu ENVCORT (kN)'],3)} kN"),
        ("Tu", f"{fmt(row['Tu diseño compat. (kN·m)'],4)} kN·m"),
        ("Armadura", f"{int(row['Nº5 sup'])}#5 sup., {int(row['Nº5 inf'])}#5 inf., {int(row['Nº5 torsión dedicadas'])}#5 torsión"),
        ("Estribo", f"{clean(row['Estribo'])}, {int(row['Ramas'])} ramas @ {int(row['s extremo DMO (mm; zona 2h)'])}/{int(row['s centro (mm)'])} mm"),
    ]
    add_table(document, ("Dato", "VC1"), inputs, widths=(6.0, 10.0), font_size=9)
    add_heading(document, "7.1 Flexión", 2)
    add_equation(document, ("d", "=", "550", "−", "40", "−", "9.5", "−", frac("15.9", "2"), "=", fmt(row["d sup (mm)"],2), " mm"))
    add_equation(document, (sub("A", "s,req"), "=", fmt(row["As req. sup (mm²)"],1), " mm²"))
    add_equation(document, (sub("A", "s,prov"), "=", "4", "·", "199", "=", fmt(row["As prov. sup (mm²)"],0), " mm²"))
    add_equation(document, (sub("φM", "n"), "−", "=", fmt(row["φMn− (kN·m)"],1), " kN·m", " ≥ ", sub("M", "u"), "−", "=", fmt(row["Mu− ENVFLEX (kN·m)"],1), " kN·m"))
    add_equation(document, (sub("φM", "n"), "+", "=", fmt(row["φMn+ (kN·m)"],1), " kN·m", " ≥ ", sub("M", "u"), "+", "=", fmt(row["Mu+ ENVFLEX (kN·m)"],1), " kN·m"))
    add_body(document, "Conclusión: el acero superior e inferior satisface la resistencia a flexión.")
    add_heading(document, "7.2 Cortante por capacidad", 2)
    add_equation(document, (sub("V", "cap"), "=", fmt(row["Vcap (kN)"],2), " kN"))
    add_equation(document, (sub("V", "u,diseño"), "=", "max", "(", fmt(row["Vu ENVCORT (kN)"],3), ",", fmt(row["Vcap (kN)"],2), ")", "=", fmt(row["Vu diseño (kN)"],2), " kN"))
    add_equation(document, (sub("φV", "n"), "=", fmt(row["φVn (kN)"],1), " kN", " ≥ ", sub("V", "u,diseño"), "=", fmt(row["Vu diseño (kN)"],1), " kN"))
    add_body(document, "En VC1 gobierna Vcap, no ENVCORT. Esto demuestra por qué no basta con leer únicamente el cortante del análisis.")
    add_heading(document, "7.3 Torsión e interacción", 2)
    add_equation(document, (sub("T", "u"), "=", fmt(row["Tu diseño compat. (kN·m)"],2), " kN·m", " > ", sub("φT", "th"), "=", fmt(row["φTth (kN·m)"],2), " kN·m"))
    add_equation(document, (frac(sub("A", "t"), "s"), "=", fmt(row["At/s (mm²/mm)"],4), " mm²/mm"))
    add_equation(document, (sub("A", "l,req"), "=", fmt(row["Al adicional req. (mm²)"],1), " mm²", " < ", sub("A", "l,prov"), "=", fmt(row["Al provisto (mm²)"],0), " mm²"))
    add_equation(document, ("Interacción V-T", "=", fmt(row["Interacción V-T"],3), " ≤ ", fmt(row["Límite V-T"],3)))
    add_callout(document, "Conclusión VC1", "flexión negativa, flexión positiva, sección a cortante, resistencia a cortante, torsión, interacción V-T y acero longitudinal torsional resultan CUMPLE.", PALE_GREEN)
    add_heading(document, "Cómo contarlo oralmente", 2)
    add_body(document, "“Primero tomo las demandas de la envolvente. Después verifico el acero longitudinal. Luego calculo los momentos probables para no subestimar el cortante. Finalmente reviso el umbral de torsión, el refuerzo transversal y longitudinal y la interacción V-T. Como todas las desigualdades son favorables, AZ devuelve CUMPLE.”")


def add_noncompliant_example(document, data):
    add_heading(document, "8. Ejemplo para explicar un NO CUMPLE: VC5, frame 84", 1)
    row = find_frame(data["rows"], 84)
    if str(row["Grupo"]) != "VC5" or str(row["ESTADO"]) != "NO CUMPLE":
        raise ValueError(
            f"El ejemplo crítico cambió: frame 84, grupo={row['Grupo']}, estado={row['ESTADO']}"
        )
    items = [
        ("Frame", row["Frame"]),
        ("Mu− / Mu+", f"{fmt(row['Mu− ENVFLEX (kN·m)'],1)} / {fmt(row['Mu+ ENVFLEX (kN·m)'],1)} kN·m"),
        ("Vu ENVCORT", f"{fmt(row['Vu ENVCORT (kN)'],1)} kN"),
        ("Tu", f"{fmt(row['Tu diseño compat. (kN·m)'],1)} kN·m"),
        ("Flexión − / +", f"{row['Flexión −']} / {row['Flexión +']}"),
        ("Cortante", row["Cortante"]),
        ("Chequeo torsión", row["Chequeo torsión"]),
        ("Interacción V-T", row["Chequeo V-T"]),
        ("Estado", row["ESTADO"]),
    ]
    add_table(document, ("Verificación", "Resultado"), items, widths=(7.0, 9.0), status_columns=(1,), font_size=9)
    add_body(document, "El frame 84 sí cumple flexión superior e inferior, pero falla cortante, el chequeo integral de torsión y la interacción V-T. Como AZ utiliza una condición AND, basta un solo incumplimiento para que el estado general sea NO CUMPLE.")
    add_callout(document, "Qué demuestra", "un resultado rojo no significa que el Excel esté mal; significa que el armado o la sección adoptada no satisface todas las verificaciones. La respuesta correcta es rediseñar y recalcular, no borrar la condición ni cambiar el texto a CUMPLE.", RED)
    add_heading(document, "Qué alternativas se revisarían", 2)
    add_bullet(document, "Aumentar dimensiones de la sección, sujeto a coordinación arquitectónica.")
    add_bullet(document, "Aumentar ramas o área del estribo y revisar la separación.")
    add_bullet(document, "Aumentar o redistribuir el acero longitudinal de torsión.")
    add_bullet(document, "Revisar el detalle cerrado del estribo y la interacción V-T.")
    add_bullet(document, "Recalcular el libro y someter el cambio a revisión profesional.")


def add_results(document, data):
    landscape_section(document)
    add_heading(document, "9. Resultados de los 12 grupos", 1)
    rows = []
    for group in GROUP_ORDER:
        row = data["summaries"][group]
        detail = data["statuses"][group]
        rows.append((
            group,
            int(row["N vigas"]),
            f"{fmt(row['Ln adopt. mín (m)'],2)}–{fmt(row['Ln adopt. máx (m)'],2)}",
            f"{int(row['b adopt. (mm)'])}×{int(row['h adopt. (mm)'])}",
            f"{fmt(row['Mu− ENVFLEX (kN·m)'],1)} / {fmt(row['Mu+ ENVFLEX (kN·m)'],1)}",
            f"{fmt(row['Vu ENVCORT (kN)'],1)} / {fmt(row['Tu diseño compat. (kN·m)'],1)}",
            f"{int(row['Nº5 sup'])}#5 / {int(row['Nº5 inf'])}#5 / {int(row['Nº5 torsión dedicadas'])}#5",
            f"{clean(row['Estribo'])}; {int(row['Ramas'])}R @{int(row['s extremo DMO (mm; zona 2h)'])}/{int(row['s centro (mm)'])}",
            f"{detail['CUMPLE']} C / {detail['NO CUMPLE']} NC",
            row["ESTADO"],
            group_failure_labels(row),
        ))
    add_table(
        document,
        ("Grupo", "N", "L (m)", "b×h", "Mu−/Mu+", "VuC/Tu", "Sup/Inf/Tor", "Estribo", "Detalle", "Resumen", "Falla resumen"),
        rows,
        widths=(1.4, 0.8, 1.8, 1.7, 2.2, 2.2, 2.5, 3.3, 1.9, 2.1, 4.4),
        status_columns=(9,),
        font_size=6.8,
    )
    portrait_section(document)
    total = data["detail_totals"]
    summary_total = data["summary_totals"]
    compliant_groups = data["summary_compliant_groups"]
    noncompliant_groups = data["summary_noncompliant_groups"]
    add_heading(document, "9.1 Dos niveles de resultado", 2)
    add_table(
        document,
        ("Nivel", "CUMPLE", "NO CUMPLE", "Interpretación"),
        (
            ("Detalle por frame", total["CUMPLE"], total["NO CUMPLE"], "Cada frame usa su demanda y su armado individual registrado."),
            ("Resumen por grupo", summary_total["CUMPLE"], summary_total["NO CUMPLE"], "Un armado común se enfrenta a la envolvente conservadora del grupo y además exige que todos sus frames cumplan."),
        ),
        widths=(4.0, 2.5, 3.0, 7.0),
        status_columns=(2,),
        font_size=8.5,
    )
    add_callout(
        document,
        "No es una contradicción",
        "VC2, VC3 y VR N1 tienen todos sus frames detallados en CUMPLE, pero su armado común de la hoja resumen no satisface la envolvente combinada. Por ejemplo, el frame 56 de VC2 usa 5#5 superiores e inferiores y 8#5 de torsión; el resumen VC2 conserva 4#5, 4#5 y 0#5, por lo que falla. Detalle y resumen responden preguntas diferentes.",
        YELLOW,
    )
    add_heading(document, f"9.2 Distribución de los {total['NO CUMPLE']} NO CUMPLE detallados", 2)
    failure_rows = []
    for group in GROUP_ORDER:
        counter = data["statuses"][group]
        if counter["NO CUMPLE"]:
            failure_rows.append((group, counter["NO CUMPLE"], ", ".join(f"{key}: {value}" for key, value in data["failures"][group].items())))
    add_table(document, ("Grupo", "Frames NC", "Verificaciones que fallan"), failure_rows, widths=(3.0, 3.0, 10.5), font_size=8.5)
    add_heading(document, "9.3 Conclusión honesta", 2)
    add_body(
        document,
        f"En la hoja detallada cumplen {total['CUMPLE']} de {sum(total.values())} frames y no cumplen {total['NO CUMPLE']}. "
        f"En el resumen común cumplen {', '.join(compliant_groups)}; los otros {summary_total['NO CUMPLE']} grupos "
        f"({', '.join(noncompliant_groups)}) requieren revisión. Ninguno de estos resultados debe presentarse como aprobación estructural.",
    )


def add_reconciliation_and_validation(document, data):
    add_heading(document, "10. Reconciliación, recálculo y alcance de esta guía", 1)
    add_body(
        document,
        "Los pasos siguientes fueron ejecutados por reconciliar_diseno_vigas.py sobre el XLSX. "
        "Este generador de la guía no vuelve a calcular el libro: abre el archivo reconciliado, audita su estructura, "
        "sus fórmulas y los valores almacenados después de ese proceso, y documenta los resultados.",
    )
    steps = [
        "Se leyó la asignación SAP y se validaron 592 frames únicos en 12 grupos.",
        "Se cruzaron frame, conectividad, longitud, sección y fuerzas ENVFLEX/ENVCORT.",
        "Se reconstruyeron las columnas A:Q de las 592 filas y los 12 resúmenes.",
        "Se creó la hoja Parametros dentro del mismo libro.",
        "Se sustituyeron referencias externas [1] por referencias internas.",
        "Las fórmulas dependientes del estribo se cambiaron para usar BA de cada fila.",
        "Se preservaron las decisiones de armado; no se forzaron estados favorables.",
        "Apache POI 5.4.1 evaluó las fórmulas y escribió sus resultados almacenados.",
        "Se validaron fórmulas, vínculos, errores, frames, grupos y estados antes de reemplazar el libro.",
    ]
    for number, text in enumerate(steps, 1):
        add_numbered(document, number, text)
    add_heading(document, "10.1 Por qué no se usó solo openpyxl", 2)
    add_body(document, "openpyxl puede leer y escribir fórmulas, pero no es un motor de cálculo de Excel. data_only=True solo devuelve el último valor almacenado. Por eso el proceso de reconciliación utilizó Apache POI para evaluar las 21.744 fórmulas y luego auditó los resultados almacenados. El generador de este Word únicamente lee y valida ese estado posterior; no invoca Apache POI.")
    add_heading(document, "10.2 Evidencias de validación", 2)
    formula_rows = [(sheet, count) for sheet, count in data["formula_by_sheet"].items()]
    formula_rows.append(("TOTAL", data["formula_count"]))
    add_table(document, ("Hoja", "Fórmulas"), formula_rows, widths=(9.0, 4.0), font_size=9)
    add_bullet(document, f"{data['formula_count']:,} fórmulas preservadas en el libro reconciliado; sus valores almacenados fueron leídos y auditados por esta guía.".replace(",", "."))
    add_bullet(document, "El recálculo con Apache POI 5.4.1 corresponde a reconciliar_diseno_vigas.py; este generador no ejecuta el motor de cálculo.")
    add_bullet(document, "Cero referencias [1] y cero partes OOXML externalLinks.")
    add_bullet(document, "Cero errores Excel almacenados (#VALUE!, #REF!, etc.).")
    add_bullet(document, "Identidad exacta de 592 frames y asignación de grupo coincidente frame a frame entre SAP y Excel.")
    add_bullet(document, f"{data['detail_totals']['CUMPLE']} CUMPLE y {data['detail_totals']['NO CUMPLE']} NO CUMPLE en detalle; son los estados almacenados del libro vigente.")
    add_callout(document, "Qué significa portable", "el libro ya no depende de una ruta OneDrive ni de un libro externo denominado [1]. La hoja Parametros viaja dentro del mismo XLSX.", PALE_GREEN)
    add_heading(document, "10.3 Entorno y archivos necesarios para reproducir la guía", 2)
    environment_rows = [(tool, value) for tool, value in data["environment"].items()]
    add_table(document, ("Componente", "Versión/uso"), environment_rows, widths=(5.0, 11.0), font_size=8.5)
    add_body(
        document,
        "La regeneración requiere este script y los archivos Diseño de vigas proyecto diseño DEF.xlsx, resultados sap.xlsx y geomatria sap.xlsx. "
        "La imagen CAD es opcional: solo se incorpora al anexo si está disponible. El recálculo del XLSX se reproduce por separado con reconciliar_diseno_vigas.py.",
    )


def add_live_demo(document):
    add_heading(document, "11. Demostración en vivo del Excel", 1)
    add_body(document, "Si durante la exposición puedes compartir el Excel, sigue este orden para no perderte:")
    demo = [
        ("1", "Parametros", "Muestra f'c, fy, φ, recubrimiento y barra #5. Señala que B9 es referencial y BA manda por fila."),
        ("2", "Todas las Vigas (592)", "Busca frame 154. Recorre A:Q y luego R:BJ hasta AZ = CUMPLE."),
        ("3", "Vigas de Carga (7)", "Muestra VC1 y explica los diferentes frames controladores de Mu−, Mu+, Vu y Tu."),
        ("4", "Todas las Vigas (592)", "Busca frame 84 y muestra que flexión cumple, pero cortante/torsión/V-T no."),
        ("5", "Vigas de Carga (7)", "Muestra VC2: detalle completo cumple, pero el armado resumen 4#5/4#5/0#5 no cubre la envolvente."),
        ("6", "Barra de fórmulas", "Selecciona AZ y explica la condición AND; selecciona BF/BG para interacción V-T."),
    ]
    add_table(document, ("Paso", "Hoja", "Acción"), demo, widths=(1.5, 5.0, 10.5), font_size=8.6)
    add_heading(document, "11.1 Ruta visual de una fila", 2)
    add_body(document, "Demanda → armado adoptado → resistencia → chequeos parciales → estado general. No saltes directamente a AZ; demuestra al menos una desigualdad de flexión y una de cortante/torsión.")
    add_heading(document, "11.2 Si una fórmula aparece como texto", 2)
    add_body(document, "Aclara que en esta guía las expresiones son ecuaciones nativas editables de Word. En Excel, la fórmula operativa aparece en la barra de fórmulas y usa referencias de celda equivalentes a la expresión matemática.")


def add_questions(document):
    add_heading(document, "12. Preguntas probables del jurado y respuestas", 1)
    qa = [
        ("¿De dónde salen los 592 frames?", "De Groups 2 - Assignments de SAP2000. Se aceptaron solo objetos Frame de los 12 grupos de vigas; no hubo duplicados."),
        ("¿Cómo se agruparon?", "Se respetaron los grupos definidos en SAP. El script no creó grupos por parecido geométrico; únicamente normalizó dos alias de nombre."),
        ("¿Todos los elementos de un grupo tienen la misma luz?", "No necesariamente. D y E muestran el rango. El grupo representa una familia SAP y el resumen usa la envolvente más desfavorable."),
        ("¿Por qué usar M3, V2 y T?", "Son los componentes locales exportados para el comportamiento principal de las vigas. La orientación debe corresponder con los ejes locales del modelo."),
        ("¿Por qué hay ENVFLEX y ENVCORT?", "ENVFLEX separa los máximos positivo/negativo de momento y aporta un cortante de flexión; ENVCORT suministra la demanda conservadora de cortante y torsión."),
        ("¿Por qué Vu diseño no siempre es Vu de SAP?", "Porque se compara con Vcap obtenido de los momentos probables. En DMO no se debe diseñar un mecanismo frágil de cortante más débil que el mecanismo probable de flexión."),
        ("¿El Excel elige automáticamente el acero?", "No completamente. Cada fila tiene nueve decisiones de armado y 36 fórmulas de comprobación. La reconciliación no cambia esas decisiones para forzar CUMPLE."),
        ("¿Por qué el diámetro de estribo no viene siempre de Parametros B9?", "Porque hay #3 y #4. BA contiene el diámetro real de cada fila y evita calcular todos los grupos con 9.5 mm."),
        ("¿Por qué VC2 detalle cumple y resumen no?", "El detalle tiene armados adaptados por frame; el resumen verifica un armado común más pequeño contra máximos que pueden venir de frames distintos."),
        ("¿Qué significa “referencia” en un estribo?", "Que ese armado aún no debe presentarse como detalle definitivo. La fórmula de cortante exige que el texto no contenga “referencia” para declarar CUMPLE."),
        ("¿Qué significa torsión despreciable?", "Que Tu no supera φTth. No significa que T sea exactamente cero; significa que no activa el diseño torsional del bloque."),
        ("¿Por qué 21.744 fórmulas?", "Hay 36 fórmulas en cada una de 592 filas detalladas y en cada una de 12 filas resumen: 36 × (592 + 12) = 21.744."),
        ("¿Cómo se recalculó el XLSX?", "Con Apache POI 5.4.1. openpyxl se usó para estructura y validación, pero no se confundió con un motor de cálculo."),
        ("¿CUMPLE significa aprobado?", "No. Es el resultado de las verificaciones implementadas. Falta revisión integral, coordinación del detalle, nudos, constructibilidad y firma profesional."),
        ("¿Qué harían con los NO CUMPLE?", "Revisar sección, acero longitudinal, ramas/diámetro/separación de estribos y torsión; luego recalcular y someter a revisión profesional."),
        ("¿El CAD dibuja los 592 frames?", "El Excel valida los 592. Las cuatro láminas muestran una viga representativa por grupo con demandas por cada tramo dibujado; no son 592 elevaciones individuales."),
    ]
    add_table(document, ("Pregunta", "Respuesta defendible"), qa, widths=(6.0, 11.0), font_size=8.2)


def add_red_flags(document):
    add_heading(document, "13. Errores que debes evitar al exponer", 1)
    errors = [
        ("No digas", "“Todo cumple”.", "Di: 538 frames detallados cumplen, 54 no; solo tres resúmenes grupales cumplen."),
        ("No digas", "“Los grupos tienen una sola luz”.", "Di: son grupos SAP y D:E conserva el rango de luces."),
        ("No digas", "“Python diseñó automáticamente las barras”.", "Di: Python reconcilió demandas; las decisiones de armado se conservaron y las fórmulas las verificaron."),
        ("No digas", "“openpyxl recalculó Excel”.", "Di: Apache POI evaluó las fórmulas; openpyxl estructuró y auditó."),
        ("No digas", "“El plano está aprobado”.", "Di: es un detalle académico, no emitir para construcción."),
        ("No ocultes", "los NO CUMPLE.", "Úsalos para demostrar que la lógica del libro detecta fallas y no maquilla resultados."),
        ("No confundas", "resumen y detalle.", "El detalle usa armado por frame; el resumen prueba un armado común contra la envolvente."),
    ]
    add_table(document, ("Advertencia", "Error", "Forma correcta"), errors, widths=(3.0, 5.0, 9.0), font_size=8.5)
    add_callout(document, "Regla de oro", "cada cifra que menciones debe poder rastrearse a una celda y a un frame. Si no puedes rastrearla, no la presentes como resultado definitivo.", RED)


def add_cheat_sheet(document, data):
    add_heading(document, "14. Hoja de repaso rápido", 1)
    total = data["detail_totals"]
    summary = data["summary_totals"]
    compliant_groups = ", ".join(data["summary_compliant_groups"])
    add_table(
        document,
        ("Dato", "Valor que debes recordar"),
        (
            ("Frames", "592 = 240 VC + 352 VR/VRAUX"),
            ("Grupos", "12 = 7 carga + 5 rigidez/auxiliar"),
            ("Secciones", "VC/VRAUX 450×550 mm; VR1/VR N1/VR2/VR3 500×550 mm"),
            ("Materiales", "f'c 28 MPa; fy = fyt 420 MPa"),
            ("Factores", "φ flexión 0.90; φ cortante/torsión 0.75"),
            ("Libro", "4 hojas; 62 columnas de diseño A:BJ"),
            ("Por fila", "17 datos + 9 decisiones + 36 fórmulas"),
            ("Fórmulas", "21.744 = 36 × (592 + 12)"),
            ("Detalle", f"{total['CUMPLE']} CUMPLE / {total['NO CUMPLE']} NO CUMPLE"),
            ("Resumen", f"{summary['CUMPLE']} CUMPLE: {compliant_groups}; {summary['NO CUMPLE']} NO CUMPLE"),
            ("Caso bueno", "VC1 / frame 154"),
            ("Caso crítico", "VC5 / frame 84"),
            ("Motor de cálculo", "Apache POI 5.4.1"),
            ("Portabilidad", "0 referencias [1], 0 externalLinks, 0 errores almacenados"),
            ("Alcance", "Académico; no emitir para construcción"),
        ),
        widths=(5.0, 11.5),
        font_size=9,
    )
    add_heading(document, "Orden mental de cualquier verificación", 2)
    add_flow(document, [
        ("DEMANDA", "Mu, Vu, Tu"),
        ("DECISIÓN", "Sección y armado"),
        ("RESISTENCIA", "φMn, φVn, At/s, Al"),
        ("ESTADO", "CUMPLE o NO CUMPLE"),
    ])
    add_heading(document, "Última respuesta si te preguntan por responsabilidad", 2)
    add_body(document, "“El Excel documenta y verifica un ejercicio académico. Los NO CUMPLE permanecen visibles y los planos indican que no deben emitirse para construcción. Una versión constructiva requiere rediseño de los incumplimientos, revisión integral y aprobación de un profesional competente.”")


def add_traceability_annex(document, data):
    landscape_section(document)
    add_heading(document, "Anexo A. Trazabilidad de fórmulas Excel", 1)
    rows = []
    descriptions = {
        "R": "As requerido superior",
        "U": "Peralte superior",
        "W": "φMn−",
        "AF": "Mpr−",
        "AH": "Vcap",
        "AI": "Vu diseño",
        "AJ": "Vc",
        "AP": "s extremo",
        "AR": "φVn",
        "AT": "φTth",
        "AV": "At/s",
        "AW": "Al requerido",
        "AZ": "Estado",
        "BF": "Interacción V-T",
        "BG": "Límite V-T",
    }
    for column in descriptions:
        formula = data["detail_formula_example"][column]
        rows.append((column, descriptions[column], formula))
    add_table(document, ("Col.", "Resultado", "Fórmula de la primera fila detallada"), rows, widths=(1.5, 4.5, 19.0), font_size=6.4)
    add_heading(document, "Anexo B. Fórmulas de estado", 1)
    add_body(document, f"Detalle AZ4: {data['detail_status_formula']}")
    add_body(document, f"Resumen AZ4: {data['summary_status_formula']}")
    add_callout(document, "Lectura", "el resumen exige sus propias verificaciones y además que no exista ningún frame NO CUMPLE dentro del grupo.", LIGHT_BLUE)
    portrait_section(document)


def add_cad_appendix(document):
    add_heading(document, "Anexo C. Relación con los planos B2", 1)
    add_body(document, "El generador CAD cruza las 592 filas del Excel con los 592 frames de SAP. Para cada grupo selecciona una cadena representativa y rotula en cada tramo dibujado Mu−, Mu+, Vu ENVFLEX, Vu ENVCORT, Tu y estado. Los paneles muestran sección, armado, estribos, estado grupal y cantidad de frames NO CUMPLE.")
    if PREVIEW.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(PREVIEW), width=Inches(6.4))
        caption = document.add_paragraph("Vista previa de las cuatro láminas B2 reconciliadas.")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.style = "Caption"
    add_callout(document, "Alcance del dibujo", "es una representación académica por grupo; no reemplaza el detalle individual de los 592 frames ni constituye aprobación para construcción.", RED)


def add_references(document):
    add_heading(document, "Referencias y archivos de respaldo", 1)
    for text in (
        "Asociación Colombiana de Ingeniería Sísmica. Reglamento Colombiano de Construcción Sismo Resistente NSR-10, especialmente el Título C.",
        "ACI 318-19, referencia de diseño de concreto configurada en el exporte del modelo SAP2000.",
        "resultados sap.xlsx — asignaciones de grupos y fuerzas por frame.",
        "geomatria sap.xlsx — conectividad, longitudes y secciones.",
        "Diseño de vigas proyecto diseño DEF.xlsx — libro reconciliado explicado en esta guía.",
        "reconciliar_diseno_vigas.py — procedimiento reproducible de actualización y recálculo.",
        "generar_planos_vigas_continuas_b2.py — generación y auditoría de las láminas B2.",
    ):
        add_bullet(document, text)
    add_callout(document, "Nota final", "la guía explica lo implementado en el libro; no sustituye una memoria profesional firmada ni una revisión completa de todos los requisitos reglamentarios y de constructibilidad.", YELLOW)


def main():
    data = load_data()
    document = Document()
    configure_document(document)
    add_cover(document)
    add_heading(document, "Tabla de contenido", 1)
    add_toc(document)
    document.add_page_break()
    add_executive_script(document, data)
    add_workflow_and_sources(document, data)
    add_grouping(document, data)
    add_architecture(document, data)
    add_column_map(document, data)
    add_formulas(document)
    add_worked_example(document, data)
    add_noncompliant_example(document, data)
    add_results(document, data)
    add_reconciliation_and_validation(document, data)
    add_live_demo(document)
    add_questions(document)
    add_red_flags(document)
    add_cheat_sheet(document, data)
    add_traceability_annex(document, data)
    add_cad_appendix(document)
    add_references(document)

    document.core_properties.title = "Guía de sustentación del Excel de diseño de vigas"
    document.core_properties.subject = "SAP, agrupación, 21.744 fórmulas, resultados y guion oral"
    document.core_properties.author = "Grupo 6 — Edificio residencial Santa Marta"
    document.core_properties.comments = "Entrega académica. No emitir para construcción."
    document.save(OUTPUT)
    print(f"OK {OUTPUT}")
    print(f"Frames={EXPECTED_FRAMES} | grupos={EXPECTED_GROUPS} | fórmulas={data['formula_count']}")
    print("Incluye ecuaciones OMML nativas, ejemplo CUMPLE/NO CUMPLE, guion y preguntas.")


if __name__ == "__main__":
    main()
