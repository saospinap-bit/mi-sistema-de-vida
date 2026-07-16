# -*- coding: utf-8 -*-
"""Genera el documento Word: explicacion detallada + verificacion y correcciones."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IMG = "_img"
AZUL = RGBColor(0x1F, 0x3B, 0x73)
GRIS = RGBColor(0x40, 0x40, 0x40)
VERDE = RGBColor(0x00, 0x66, 0x00)
ROJO = RGBColor(0xB0, 0x1C, 0x1C)
NARANJA = RGBColor(0xB0, 0x60, 0x00)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11)

def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs: r.font.color.rgb = AZUL
    return p

def par(text, size=11, bold=False, italic=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_img(name, width_cm=15, caption=None):
    path = os.path.join(IMG, name)
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(width_cm))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            par(caption, size=9, italic=True, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    else:
        par(f"[Falta imagen: {name}]", italic=True, color=ROJO)

# ===================== PORTADA =====================
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("MEMORIA DE CÁLCULO DE CANTIDADES DE OBRA")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = AZUL
par("Explicación detallada, verificación contra planos y correcciones",
    size=13, italic=True, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
par("PROYECTO: Módulo 1 – Colegio", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
par("Fundamentos de Construcción – Universidad Nacional de Colombia", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
par("Etapa: Obra gris  |  Grupo 3  |  Norma NSR-10  |  Pórticos de concreto reforzado",
    size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = p.add_run("Cada cantidad se explica indicando de qué plano y de qué cota sale,\ncon la fórmula, el resultado y la verificación realizada contra los planos.")
rr.italic = True; rr.font.size = Pt(10.5); rr.font.color.rgb = GRIS
doc.add_page_break()

# ===================== 1. OBJETIVO =====================
h("1. Objetivo y alcance", 1)
par("Este documento explica, de forma trazable, cómo se calcularon las cantidades de obra del "
    "proyecto (archivo MEMORIAS DE CALCULO.xlsx) y documenta la verificación de cada dimensión "
    "contra los planos entregados. Para cada actividad se indica:")
bullet("qué se mide y en qué unidad;", "Actividad e ítem: ")
bullet("de qué plano y de qué cota sale cada dimensión;", "Origen del valor: ")
bullet("la fórmula geométrica y el resultado;", "Cálculo: ")
bullet("qué imagen/plano pegar en la casilla «IMAGEN/REFERENCIA» del formato.", "Foto de referencia: ")

# ===================== 2. METODOLOGIA =====================
h("2. Metodología general de cálculo", 1)
par("Las cantidades de obra son la medición de cada actividad en su unidad de pago. El formato "
    "usa cuatro datos por renglón: Ancho, Alto (o espesor), Largo y Cantidad. El total del "
    "renglón es el producto de esos cuatro campos:")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = p.add_run("TOTAL = Ancho × Alto × Largo × Cantidad")
rr.bold = True; rr.font.size = Pt(13); rr.font.color.rgb = AZUL
par("Según la unidad, algún campo se deja en 1 para que la fórmula entregue la unidad correcta:", space_after=4)
tb = doc.add_table(rows=1, cols=3); tb.style = "Light Grid Accent 1"; tb.alignment = WD_TABLE_ALIGNMENT.CENTER
hh = tb.rows[0].cells
hh[0].text="Unidad"; hh[1].text="Qué representa"; hh[2].text="Cómo se arma"
for a,b,c in [
    ("M3","Concreto, excavaciones, rellenos","Ancho × Alto × Largo × Cantidad"),
    ("M2","Placas, muros, geotextil, pañete, adoquín","un campo = 1 (Área × Cantidad o Ancho × Largo)"),
    ("ML","Columnetas de confinamiento","Largo × Cantidad (Ancho = 1)"),
    ("KG","Acero de refuerzo","Longitud de varillas × peso lineal (kg/m)"),
]:
    row=tb.add_row().cells; row[0].text=a; row[1].text=b; row[2].text=c
    for cc in row: cc.paragraphs[0].runs[0].font.size = Pt(10)
par("")
par("Las dimensiones se leen de los planos; por eso el formato pide la «IMAGEN/REFERENCIA» "
    "(el pantallazo del plano donde se ve la cota usada).", space_after=8)

par("2.1  Cómo se calcula el ACERO (por peso)", bold=True, color=AZUL, space_after=2)
par("El acero se paga por kilogramos. El método es: para cada elemento se leen del DESPIECE "
    "el número, el calibre y la longitud de cada barra; se suma la longitud total de cada "
    "calibre y se multiplica por su masa lineal (kg/m). Al final se añade 10 % por desperdicio "
    "y traslapos:", space_after=4)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
rr=p.add_run("kg = Σ [ Longitud total de cada calibre (m) × masa (kg/m) ] × 1.10")
rr.bold=True; rr.font.size=Pt(12.5); rr.font.color.rgb=AZUL
par("Masa lineal de la varilla (tabla de cantidad de varillas por tonelada):", space_after=4)
tb = doc.add_table(rows=1, cols=4); tb.style="Light Grid Accent 1"; tb.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,tx in enumerate(["Barra","Pulg.","mm","Masa (kg/m)"]): tb.rows[0].cells[i].text=tx
for a,b,c,d in [
    ("N°2","1/4\"","6.4","0.250"),("N°3","3/8\"","9.5","0.560"),("N°4","1/2\"","12.7","0.994"),
    ("N°5","5/8\"","15.9","1.552"),("N°6","3/4\"","19.1","2.235"),("N°7","7/8\"","22.2","3.042"),
    ("N°8","1\"","25.4","3.973"),
]:
    row=tb.add_row().cells; row[0].text=a; row[1].text=b; row[2].text=c; row[3].text=d
    for cc in row:
        for rn in cc.paragraphs[0].runs: rn.font.size=Pt(9.5)
par("Ejemplo (columnas Tipo 1): 4 columnas × 6 barras N°6 × 5.85 m = 140.4 m de N°6; "
    "140.4 × 2.235 kg/m = 313.8 kg solo en el longitudinal de ese tipo.",
    size=10, italic=True, color=GRIS, space_after=6)

# ===================== 3. PLANOS =====================
h("3. Planos de soporte", 1)
add_img("plano_cimentacion.png", 16, "Figura 3.1. Cimentación: ejes 6.00+6.60+5.80+1.40 = 19.80 m; transversal 7.30/7.70 m; vigas VC 0.40×0.45; zapatas Z-1 (2.30×2.30) y Z-2 (2.15×2.15); placa de piso e=0.10.")
add_img("plano_columnas.png", 16, "Figura 3.2. Columnas: dos tipos de columna (ver detalles ampliados en la Sección 4).")
add_img("plano_cubierta.png", 16, "Figura 3.3. Placa de cubierta: losa aligerada e=0.08 en N+4.90; vigas VG 0.40×0.50; viguetas 0.20×0.40.")
add_img("plano_detalles_mamp.png", 16, "Figura 3.4. Detalles de mampostería confinada.")
add_img("arq_full_hi.png", 16, "Figura 3.5. Planta arquitectónica: ejes 1–4 y A–B y áreas de cada espacio (NO trae cotas totales).")
doc.add_page_break()

# ===================== 4. VERIFICACION Y CORRECCIONES =====================
h("4. Verificación contra planos y correcciones aplicadas", 1)
par("Se revisó cada dimensión de la memoria contra los planos. A continuación el resultado, "
    "separado en lo verificado, lo corregido y lo que debe confirmarse en el DWG.", space_after=8)

par("4.1  Valores VERIFICADOS (coinciden con el plano)", bold=True, color=VERDE, space_after=2)
tb = doc.add_table(rows=1, cols=3); tb.style="Light Grid Accent 1"
for i,tx in enumerate(["Dato","Memoria","Plano"]): tb.rows[0].cells[i].text=tx
for a,b,c in [
    ("Longitud entre ejes 1–4","19.80 m","6.00+6.60+5.80+1.40 = 19.80"),
    ("Luz transversal A–B","7.30 m","7.30 (7.70 total)"),
    ("Vigas de cimentación","0.40×0.45; 2 long + 5 transv","VC-A/VC-B y VC-1..4 = 0.40×0.45"),
    ("Placa de contrapiso","e = 0.10","«PLACA PISO e:0.10m»"),
    ("Vigas de cubierta","0.40×0.50","VG-6/VG-7 = 40×50"),
    ("Placa de cubierta","losa aligerada N+4.90","«LOSA ALIGERADA N+4.90 e:0.08»"),
]:
    row=tb.add_row().cells; row[0].text=a; row[1].text=b; row[2].text=c
    for cc in row:
        for rn in cc.paragraphs[0].runs: rn.font.size=Pt(9.5)
par("")

par("4.2  Valores CORREGIDOS (no coincidían con el plano)", bold=True, color=ROJO, space_after=2)
tb = doc.add_table(rows=1, cols=3); tb.style="Light Grid Accent 1"
for i,tx in enumerate(["Ítem","Antes (Excel)","Ahora (según plano)"]): tb.rows[0].cells[i].text=tx
for a,b,c in [
    ("2.1.1 Excavación masiva","prof. 0.55 m","0.40 m (sub-base 0.30 + placa 0.10; el recebo va aparte)"),
    ("2.1.2 Sub-base","e = 0.15 m","e = 0.30 m (el detalle dice 'SUB-BASE 30 cm')"),
    ("2.1.3 Excavación manual","solo zanjas 0.40 prof.","zanjas 0.45 + fosos de 8 zapatas hasta −1.55"),
    ("2.2.2 Solado","solo bajo vigas (38.05)","+ solado bajo 8 zapatas = 77.70 m²"),
    ("2.2.3 Zapatas concreto","no existía","AGREGADO: 4×Z-1 + 4×Z-2 = 15.86 m³"),
    ("2.3.2 Acero zapatas","no existía","AGREGADO (estimación): ≈668.6 kg"),
    ("4.1.1 Columnas","10 col de 0.40×0.40, h=4.90","8 col en 2 tipos: 4×(0.40×0.40, h=4.75) + 4×(0.50×0.50, h=4.95)"),
    ("4.3.3.1 Placa cubierta","M2 pero fórmula con espesor","Área de losa = 7.30 × 18.52 (Alto=1)"),
    ("4.3.3.1 Viguetas","0.20 × 0.32","0.20 × 0.40 (el plano dice 20×40)"),
    ("4.7.1 Acero columnas","10 col × (2#6+2#7)","Tipo1: 6#6 ; Tipo2: 4#6+4#7 (kg desde cartilla)"),
]:
    row=tb.add_row().cells; row[0].text=a; row[1].text=b; row[2].text=c
    for cc in row:
        for rn in cc.paragraphs[0].runs: rn.font.size=Pt(9.5)
par("")

par("4.3  Valores POR CONFIRMAR en el DWG (no acotados en los PDF)", bold=True, color=NARANJA, space_after=2)
bullet("no está acotada en los planos PDF; la retícula estructural acotada es 19.80 × 7.30/7.70. Confirmar en el DWG/arquitectónico y pegar el pantallazo con la cota. Afecta: excavación, sub-base, recebo, geotextil y placa de contrapiso.", "Huella 9.85 × 21.95 m: ")
bullet("suma de longitudes de todos los muros; medir sobre el plano de mampostería y adjuntar el pantallazo.", "L = 173.60 m de muro (5.1.1): ")
bullet("una cada ~3 m; contar sobre el plano.", "58 columnetas (5.1.2): ")
bullet("zonas exteriores; medir sobre el plano de adoquín/arquitectónico.", "9 áreas de adoquín (8.1.1): ")
par("")

par("4.4  Otras observaciones técnicas", bold=True, color=AZUL, space_after=2)
bullet("Los planos especifican concreto de 4000 psi (280 kg/cm²), mientras que el presupuesto y la memoria usan 3000 psi. Verificar con el profesor cuál rige.")
bullet("El plano tiene 8 zapatas aisladas (4 Z-1 2.30×2.30×0.40 bajo las columnas de 0.50, y 4 Z-2 2.15×2.15×0.40 bajo las de 0.40). YA SE AGREGARON a la memoria (ítems 2.2.3 concreto y 2.3.2 acero); falta crear esas dos líneas en el presupuesto.")
par("")

par("Evidencia de las correcciones (recortes ampliados de los planos):", bold=True, space_after=4)
add_img("crop_col_det1.png", 8.0, "Detalle Columna Tipo 1: B=0.40 H=0.40, h=4.70, refuerzo 6#6.")
add_img("crop_col_det2.png", 8.0, "Detalle Columna Tipo 2: B=0.50 H=0.50, h=4.90, refuerzo 4#6+4#7.")
add_img("crop_cub_seccion.png", 13, "Detalle de la losa aligerada: viguetas 0.20×0.40 @1.00 m con malla electrosoldada.")
doc.add_page_break()

# ===================== 5. CALCULO DETALLADO =====================
h("5. Cálculo detallado por actividad", 1)
par("Para cada ítem: origen de las dimensiones, fórmula y resultado (ya con las correcciones).")

def item_block(cap, code, title, unit, plano_caption, dims, formula, resultado, obs=None, scheme=None, scheme_cap=None):
    h(f"Ítem {code} — {title}  ({unit})", 2)
    par(f"Capítulo: {cap}", size=10, italic=True, color=GRIS, space_after=4)
    par("Origen de cada valor (de dónde sale):", bold=True, space_after=2)
    tb = doc.add_table(rows=1, cols=2); tb.style = "Light List Accent 1"
    tb.rows[0].cells[0].text="Dimensión / dato"; tb.rows[0].cells[1].text="De dónde se obtiene"
    for d,s in dims:
        row=tb.add_row().cells; row[0].text=d; row[1].text=s
        for cc in row:
            for rn in cc.paragraphs[0].runs: rn.font.size=Pt(10)
    par("")
    par("Cálculo:", bold=True, space_after=2)
    p=doc.add_paragraph(); rr=p.add_run(formula); rr.font.size=Pt(11.5); rr.bold=True; rr.font.color.rgb=AZUL
    p=doc.add_paragraph(); rr=p.add_run(f"➜ Resultado: {resultado}"); rr.bold=True; rr.font.size=Pt(11.5); rr.font.color.rgb=VERDE
    if obs: par(f"Observación: {obs}", size=10, italic=True, color=NARANJA)
    if scheme: add_img(scheme, 15, scheme_cap)
    par(f"Foto de referencia para el formato: {plano_caption}", size=10, italic=True, color=GRIS, space_after=12)

HUELLA_OBS = "La huella 9.85×21.95 no está acotada en los PDF; confirmar en el DWG y pegar el pantallazo de la cota."

h("Capítulo 2 – Cimentación", 2)
par("Todo el capítulo 2 se apoya en un único sistema de niveles. El siguiente corte muestra "
    "cómo encajan las capas (el fondo de cada una es el techo de la siguiente), verificado con "
    "los tres detalles del plano: losa de contrapiso, planta de cimentación (NE −1.10) y detalle "
    "de zapata (desplante −1.50).", space_after=4)
add_img("esq_seccion_general.png", 16, "Corte lógico de la cimentación con niveles reales (desde N ±0.00).")
par("")
item_block("2. Cimentación","2.1.1","Excavación a máquina conglomerado 0–2 m (incl. retiro)","M3",
    "Figura 3.1 + detalle de losa de contrapiso.",
    [("Ancho = 9.85 m / Largo = 21.95 m","Huella del módulo (POR CONFIRMAR en DWG)."),
     ("Alto = 0.40 m","= sub-base 0.30 + placa 0.10 (el paquete de piso que muestra el detalle de la losa). El recebo NO se suma aquí: es un relleno que va en otra parte (backfill sobre las vigas).")],
    "V = 9.85 × 0.40 × 21.95 × 1","86.48 m³",
    obs="CORREGIDO: 0.60 → 0.40 m. El detalle de la losa solo muestra sub-base 0.30 + placa 0.10. El recebo (2.1.4) es relleno aparte, no capa bajo la placa.")
item_block("2. Cimentación","2.1.2","Relleno granular sub-base B-400","M3",
    "Figura 3.1 + detalle de losa de contrapiso.",
    [("Ancho = 9.85 / Largo = 21.95","Huella del módulo (POR CONFIRMAR)."),
     ("Alto = 0.30 m","Espesor de la sub-base. El detalle dice 'SUB-BASE DE 30 cm' y 'Mínimo 0.30m'.")],
    "V = 9.85 × 0.30 × 21.95 × 1","64.86 m³",
    obs="CORREGIDO: 0.15 → 0.30 m (lo dice el detalle de la losa).")
item_block("2. Cimentación","2.1.3","Excavación manual material común 0–2 m","M3",
    "Figura 3.1 (vigas VC y zapatas) + detalle de zapata.",
    [("Zanjas de vigas: 0.40×0.45×19.80 (×2) y ×7.30 (×5)","Zanjas para las vigas de cimentación, de −0.65 a −1.10 (0.45 = alto de la viga)."),
     ("Fosos Z-1: 2.30×2.30, prof 0.90, ×4","Zapatas grandes, de −0.65 a −1.55 (desplante)."),
     ("Fosos Z-2: 2.15×2.15, prof 0.90, ×4","Zapatas pequeñas, misma profundidad.")],
    "V = 7.13 + 6.57 (zanjas) + 19.04 + 16.64 (zapatas)","49.38 m³",
    obs="CORREGIDO: la excavación manual va por debajo del nivel masivo (−0.65). Incluye las zanjas de vigas (0.45 m) y los fosos de las 8 zapatas hasta el desplante (−1.55). Estimación sin descontar solapes (conservador).")
item_block("2. Cimentación","2.1.4","Relleno material común (recebo)","M3",
    "Figura 3.1 + Esquema 1.",
    [("Ancho = 9.85 / Largo = 21.95","Huella (POR CONFIRMAR)."),
     ("Alto = 0.20 m","Espesor de recebo de nivelación.")],
    "V = 9.85 × 0.20 × 21.95 × 1","43.24 m³", obs=HUELLA_OBS)
item_block("2. Cimentación","2.1.6","Geotextil NT 1600","M2",
    "Figura 3.1 + Esquema 1.",
    [("Ancho = 9.85 / Largo = 21.95","Huella (POR CONFIRMAR)."),
     ("Alto = 1","Se mide por área.")],
    "A = 9.85 × 1 × 21.95 × 1","216.21 m²", obs=HUELLA_OBS)
item_block("2. Cimentación","2.2.1","Vigas de cimentación concreto 3000 psi","M3",
    "Figura 3.1 (rótulo B=0.40 H=0.45).",
    [("Sección 0.40 × 0.45","Rótulo de las vigas VC en el plano."),
     ("Long. 19.80 × 2","VC-A y VC-B."),
     ("Transv. 7.30 × 5","VC-1..4 (5 ejes).")],
    "V = (0.40×0.45×19.80×2)+(0.40×0.45×7.30×5) = 7.13 + 6.57","13.70 m³",
    scheme="esq_viga_cim.png", scheme_cap="Esquema 3. Sección de viga de cimentación 0.40×0.45.")
item_block("2. Cimentación","2.2.2","Concreto 1500 psi solado e=0.05","M2",
    "Figura 3.1 + detalle de zapata (solado de limpieza).",
    [("Bajo vigas: 0.50 × long. (19.80×2 y 7.30×5)","Solado 10 cm más ancho que la viga."),
     ("Bajo zapatas: Z-1 2.30×2.30 ×4 y Z-2 2.15×2.15 ×4","El detalle de zapata muestra solado de limpieza bajo cada zapata."),
     ("Alto = 1","Se mide por área.")],
    "A = 38.05 (vigas) + 39.65 (zapatas)","77.70 m²",
    obs="AGREGADO: solado también bajo las 8 zapatas (antes solo bajo vigas).")
item_block("2. Cimentación","2.2.3","Zapatas concreto 3000 psi","M3",
    "Figura 3.1 (cuadro de zapatas) + detalle de zapata.",
    [("Z-1: 2.30×2.30×0.40, ×4","Zapatas bajo columnas 0.50×0.50 (ejes 2A,3A,2B,3B)."),
     ("Z-2: 2.15×2.15×0.40, ×4","Zapatas bajo columnas 0.40×0.40 (ejes 1A,4A,1B,4B)."),
     ("H = 0.40 m","Altura de zapata del cuadro de zapatas.")],
    "V = (2.30²×0.40×4) + (2.15²×0.40×4) = 8.46 + 7.40","15.86 m³",
    obs="AGREGADO: las zapatas estaban en el plano pero no en la memoria. Falta agregar la línea 2.2.3 al presupuesto.")
item_block("2. Refuerzos / Acero","2.3.1","Acero figurado vigas de cimentación","KG",
    "Figura 3.1 (despiece de vigas VC-1..4, 3', A, B).",
    [("Longitudinal #6 = 236.7 m","3#6 por viga (VC-1..4, 3') + VC-A/B: 3#6 L=7.60 y 7.40; L de cada barra del despiece."),
     ("Longitudinal #5 = 276.2 m","3#5 por viga transversal + VC-A/B: 3#5, 2#5, 1#5 (bastones); L del despiece."),
     ("Estribos #3 = 565.6 m","Transversales: 5×36 estribos L=1.58; longitudinales: 2×89 estribos L=1.58."),
     ("Peso lineal (kg/m)","#6=2.235 ; #5=1.552 ; #3=0.560 (tabla de masas de varilla)."),
     ("+10 %","Desperdicio y traslapos.")],
    "Long: 236.7×2.235 + 276.2×1.552 = 957.7 kg ; Estribos: 565.6×0.560 = 316.8 kg ; (957.7+316.8) × 1.10","1 401.9 kg",
    obs="RECALCULADO por peso con la tabla de masas (L de cada calibre × kg/m + 10%). VC-3 se asume igual a VC-2.")
item_block("2. Refuerzos / Acero","2.3.2","Acero figurado zapatas","KG",
    "Figura 3.1 (cuadro de zapatas: #5 @0.20 ambos sentidos).",
    [("Z-1 (×4) 2.30×2.30","Parrilla #5 @0.20 en 2 sentidos. n=lado/0.20+1=12 barras/sentido; L=lado−0.15=2.15 m."),
     ("Z-2 (×4) 2.15×2.15","Parrilla #5 @0.20 en 2 sentidos. n=11 barras/sentido; L=2.00 m."),
     ("Peso lineal #5 = 1.552 kg/m","Tabla de masas de varilla."),
     ("+10 %","Desperdicio y traslapos.")],
    "382.4 m de #5 × 1.552 = 593.5 kg × 1.10","652.8 kg",
    obs="RECALCULADO por peso con la tabla de masas. Confirmar lados exactos de zapata en el DWG. Falta la línea 2.3.2 en el presupuesto.")

doc.add_page_break()
h("Capítulo 4 – Estructura en concreto", 2)
item_block("4. Estructura","4.1.1","Columnas concreto 3000 psi","M3",
    "Figura 3.2 y recortes del despiece (Sección 4).",
    [("Tipo 1: 0.40×0.40, h=4.75, 4 und","Ejes 1A, 1B, 3'B, 4A. Altura de −0.05 a +4.70 = 4.75 m (niveles del despiece)."),
     ("Tipo 2: 0.50×0.50, h=4.95, 4 und","Ejes 2A, 2B, 3A, 3B. Altura de −0.05 a +4.90 = 4.95 m."),
     ("Total 8 columnas","Contadas y verificadas en el despiece ('Son 4' + 'Son 4'). NO son 10.")],
    "V = (0.40×0.40×4.75×4) + (0.50×0.50×4.95×4) = 3.04 + 4.95","7.99 m³",
    obs="CORREGIDO: antes 10 columnas iguales de 0.40×0.40 h=4.90. El plano muestra 8 columnas en dos tipos y las alturas se leen de −0.05 a +4.70/+4.90.",
    scheme="esq_columna2.png", scheme_cap="Esquema 4. Columnas: 8 unidades en dos tipos.")
item_block("4. Estructura","4.2.1","Concreto 3000 psi viga cubierta / cinta","M3",
    "Figura 3.3 (VG rotuladas B=0.40 H=0.50).",
    [("Sección 0.40 × 0.50","Rótulo de las vigas de cubierta."),
     ("Long. 19.80 × 2","Vigas sobre ejes A y B."),
     ("Transv. 7.30 × 5","5 ejes.")],
    "V = (0.40×0.50×19.80×2)+(0.40×0.50×7.30×5) = 7.92 + 7.30","15.22 m³")
item_block("4. Estructura","4.3.1","Placa de concreto contrapiso","M3",
    "Figura 3.1 (placa e=0.10) + Esquema 1.",
    [("Ancho = 9.85 / Largo = 21.95","Huella (POR CONFIRMAR)."),
     ("Alto = 0.10 m","Espesor de la placa (rótulo del plano).")],
    "V = 9.85 × 0.10 × 21.95 × 1","21.62 m³", obs=HUELLA_OBS)
item_block("4. Estructura","4.3.3.1","Placa de cubierta aligerada","M2",
    "Figura 3.3 + recorte de sección (Sección 4) + Esquema 5.",
    [("Ancho = 7.30 m","Luz transversal."),
     ("Largo = 5.55+6.10+6.87 = 18.52 m","Desarrollo de la cubierta."),
     ("Alto = 1","Item por M2: se toma el ÁREA de losa."),
     ("Viguetas 0.20 × 0.40","Incluidas en el m² de losa aligerada (el concreto va en el APU).")],
    "A = 7.30 × 1 × 18.52 × 1","135.20 m²",
    obs="CORREGIDO: item por M2 → área de losa (antes multiplicaba por el espesor). Viguetas 0.20×0.40 (no 0.20×0.32).",
    scheme="esq_cubierta.png", scheme_cap="Esquema 5. Placa de cubierta aligerada y viguetas.")
item_block("4. Refuerzos / Acero","4.7.1","Acero figurado columnas","KG",
    "Figura 3.2 y recortes del despiece (crop_col_det1 y det2).",
    [("Tipo 1 (×4): 6#6 L=5.85","Longitudinal 0.40×0.40 → 4×6×5.85 = 140.4 m de #6."),
     ("Tipo 2 (×4): 4#6+4#7 L=5.95","Longitudinal 0.50×0.50 → 95.2 m de #6 y 95.2 m de #7."),
     ("Estribos #3 = 762.7 m","T1: 42 est. L=1.48 + 42 grapas L=0.54; T2: 42 est. L=1.88 + 42 grapas L=0.64; ×4 c/tipo."),
     ("Peso lineal (kg/m)","#6=2.235 ; #7=3.042 ; #3=0.560 (tabla de masas)."),
     ("+10 %","Desperdicio y traslapos.")],
    "Long: 235.6×2.235 + 95.2×3.042 = 816.2 kg ; Estribos: 762.7×0.560 = 427.1 kg ; (816.2+427.1) × 1.10","1 367.6 kg",
    obs="RECALCULADO por peso con la tabla de masas. Antes era una estimación (≈791.8 kg); ahora se cuenta cada barra del despiece (6#6 en T1; 4#6+4#7 en T2) más estribos y grapas.")
item_block("4. Refuerzos / Acero","4.8.1","Acero figurado vigas de cubierta","KG",
    "Figura 3.3 (despiece VG-1..4 y VG-6/VG-7).",
    [("Longitudinal #5 = 252.8 m","Bastones y corridos #5 de VG-1..4 (2#5) y VG-6/7 (3#5); L del despiece."),
     ("Longitudinal #6 = 196.6 m","2#6 en VG-1..4 y 3#6 en VG-6/7; L del despiece."),
     ("Estribos #3 = 1 102.1 m","Transversales: 4×73 est. L=1.68; longitudinales: 2×182 est. L=1.68."),
     ("Peso lineal (kg/m)","#5=1.552 ; #6=2.235 ; #3=0.560 (tabla de masas)."),
     ("+10 %","Desperdicio y traslapos.")],
    "Long: 252.8×1.552 + 196.6×2.235 = 831.7 kg ; Estribos: 1102.1×0.560 = 617.2 kg ; (831.7+617.2) × 1.10","1 593.8 kg",
    obs="RECALCULADO por peso con la tabla de masas. NO incluye las viguetas de la placa aligerada (6 uds 0.20×0.40 con #4 y estribos #3 ≈ 1 689 kg): esas se contabilizan con la placa de cubierta si el APU lo pide.")

doc.add_page_break()
h("Capítulo 5 – Mampostería", 2)
item_block("5. Mampostería","5.1.1","Muros bloque hueco No. 4","M2",
    "Figura 3.4 / plano de mampostería.",
    [("Longitud total = 173.60 m","Suma de longitudes de todos los muros (POR CONFIRMAR en DWG)."),
     ("Alto = 2.90 m","Altura libre piso a viga de cubierta."),
     ("Descuento de vanos = 20 %","Puertas y ventanas.")],
    "A = (173.60×1×2.90×1) − (173.60×1×2.90×0.20) = 503.44 − 100.69","402.75 m²",
    obs="Confirmar los 173.60 m midiendo los muros en el plano y adjuntar el pantallazo.",
    scheme="esq_muro.png", scheme_cap="Esquema 6. Elevación de muro (área bruta − 20 % vanos).")
item_block("5. Mampostería","5.1.2","Columneta de confinamiento No. 5","ML",
    "Figura 3.4.",
    [("Cantidad = 58","Una columneta cada ~3 m (POR CONFIRMAR contando en el plano)."),
     ("Alto = 2.90 m","Igual a la altura del muro."),
     ("Ancho = 1","Se mide por longitud (ml).")],
    "L = 1 × 1 × 2.90 × 58","168.20 ml",
    obs="Confirmar el número de columnetas contándolas en el plano de mampostería.")
item_block("5. Mampostería","5.1.4","Pañete liso impermeabilizado 1:4","M2",
    "Figura 3.4.",
    [("Área de muro = 402.75 m²","Tomada del ítem 5.1.1."),
     ("Factor = 65 %","Porcentaje del área de muro que lleva pañete.")],
    "A = 402.75 × 1 × 0.65 × 1","261.79 m²")
item_block("5. Refuerzos / Acero","5.2.1","Acero mampostería confinada","KG",
    "Figura 3.4 (detalle de columneta y viga cinta V2).",
    [("Columnetas 2#4 (0.20×0.15)","Del detalle de columneta: 2 N°4 verticales + estribo N°3 c/0.20. Estimado 24 uds h=2.70."),
     ("Cinta V2: 2#3 + gancho N°2 (1/4\") c/0.30","Del detalle de viga cinta. Estimado 70 m de cinta."),
     ("Peso lineal (kg/m)","#4=0.994 ; #3=0.560 ; #2=0.250 (tabla de masas)."),
     ("+10 %","Desperdicio.")],
    "Long: 207.2 kg ; Estribos+ganchos: 133.3 kg ; (207.2+133.3) × 1.10","374.6 kg (estimación)",
    obs="ESTIMACIÓN por peso con la tabla de masas usando el detalle real (columneta 2#4+est.#3; cinta 2#3+gancho#2). Las cantidades (n° de columnetas y metros de cinta) deben confirmarse en el plano de mampostería FC3.")

doc.add_page_break()
h("Capítulo 8 – Pisos", 2)
item_block("8. Pisos","8.1.1","Adoquín en concreto 0.20×0.10×0.06 m (suministro e instalación)","M2",
    "Figura 3.5 / plano de adoquín.",
    [("9 áreas parciales (m²)","18.62+17.49+22.25+7.76+14.60+15.00+1.91+1.91+15.76 (POR CONFIRMAR midiendo en el plano)."),
     ("Cantidad = 1 c/u","Cada área una vez.")],
    "A = 18.62+17.49+22.25+7.76+14.60+15.00+1.91+1.91+15.76","115.30 m²",
    obs="El plano de adoquín entregado en PDF salió sin cotas; medir las zonas exteriores en el DWG/arquitectónico y adjuntar el pantallazo.")

# ===================== 6. RESUMEN =====================
doc.add_page_break()
h("6. Resumen de cantidades de obra (corregido)", 1)
resumen = [
    ("2.1.1","Excavación a máquina 0–2 m","M3","86.48","CORREGIDO / huella"),
    ("2.1.2","Relleno sub-base B-400","M3","64.86","CORREGIDO / huella"),
    ("2.1.3","Excavación manual (zanjas + zapatas)","M3","49.38","CORREGIDO"),
    ("2.1.4","Relleno común (recebo)","M3","43.24","por confirmar huella"),
    ("2.1.6","Geotextil NT 1600","M2","216.21","por confirmar huella"),
    ("2.2.1","Vigas de cimentación 3000 psi","M3","13.70","verificado"),
    ("2.2.2","Solado 1500 psi e=0.05","M2","77.70","CORREGIDO (+ zapatas)"),
    ("2.2.3","Zapatas concreto 3000 psi","M3","15.86","AGREGADO"),
    ("2.3.1","Acero vigas de cimentación","KG","1 401.9","RECALCULADO (tabla masas)"),
    ("2.3.2","Acero zapatas","KG","652.8","AGREGADO / RECALCULADO"),
    ("4.1.1","Columnas 3000 psi (8 und, 2 tipos)","M3","7.99","CORREGIDO"),
    ("4.2.1","Vigas de cubierta 3000 psi","M3","15.22","verificado"),
    ("4.3.1","Placa de contrapiso","M3","21.62","por confirmar huella"),
    ("4.3.3.1","Placa de cubierta aligerada","M2","135.20","CORREGIDO"),
    ("4.7.1","Acero columnas","KG","1 367.6","RECALCULADO (tabla masas)"),
    ("4.8.1","Acero vigas de cubierta","KG","1 593.8","RECALCULADO (tabla masas)"),
    ("5.1.1","Muros bloque No.4","M2","402.75","por confirmar L"),
    ("5.1.2","Columneta de confinamiento","ML","168.20","por confirmar cant."),
    ("5.1.4","Pañete liso impermeabilizado","M2","261.79","verificado"),
    ("5.2.1","Acero mampostería confinada","KG","374.6","RECALCULADO (estimación)"),
    ("8.1.1","Adoquín en concreto","M2","115.30","por confirmar áreas"),
]
tb = doc.add_table(rows=1, cols=5); tb.style="Medium Shading 1 Accent 1"
for i,tx in enumerate(["Ítem","Actividad","Und","Cantidad","Estado"]): tb.rows[0].cells[i].text=tx
for it,de,un,ca,es in resumen:
    row=tb.add_row().cells; row[0].text=it; row[1].text=de; row[2].text=un; row[3].text=ca; row[4].text=es
    for cc in row:
        for rn in cc.paragraphs[0].runs: rn.font.size=Pt(9)

# ===================== 7. IMAGEN/REFERENCIA =====================
doc.add_page_break()
h("7. Cómo llenar la columna «IMAGEN/REFERENCIA»", 1)
par("En cada renglón del formato hay una casilla «IMAGEN/REFERENCIA». Ahí va la «foto de dónde "
    "salió el dato». Procedimiento:")
bullet("Abre el plano PDF del ítem (la figura indicada en cada bloque).")
bullet("Recorta la zona donde se ve la cota usada (rótulo de la viga, altura de columna, luz entre ejes, etc.).")
bullet("Pega el recorte en la casilla IMAGEN/REFERENCIA de ese ítem en el Excel.")
par("Correspondencia ítem → plano:", bold=True, space_after=2)
tb = doc.add_table(rows=1, cols=2); tb.style="Light Grid Accent 1"
tb.rows[0].cells[0].text="Ítems"; tb.rows[0].cells[1].text="Plano a recortar"
for a,b in [
    ("2.1.1 / 2.1.2 / 2.1.4 / 2.1.6 / 4.3.1","Cimentación + corte de piso (Fig 3.1 / Esquema 1)"),
    ("2.1.3 / 2.2.1 / 2.2.2 / 2.3.1","Cimentación – vigas VC y despiece (Fig 3.1)"),
    ("4.1.1 / 4.7.1","Detalle de columnas (Fig 3.2 y recortes Sección 4)"),
    ("4.2.1 / 4.3.3.1 / 4.8.1","Placa de cubierta y vigas VG (Fig 3.3)"),
    ("5.1.1 / 5.1.2 / 5.1.4 / 5.2.1","Detalles de mampostería (Fig 3.4)"),
    ("8.1.1","Adoquín / arquitectónico con áreas (Fig 3.5)"),
]:
    row=tb.add_row().cells; row[0].text=a; row[1].text=b
    for cc in row:
        for rn in cc.paragraphs[0].runs: rn.font.size=Pt(10)
par("")
par("Nota: las correcciones de columnas y viguetas ya quedaron aplicadas en MEMORIAS DE "
    "CALCULO.xlsx. Los valores marcados «por confirmar» dependen de una cota que no está en los "
    "PDF; confírmalos en el DWG y quedan cerrados.", size=10, italic=True, color=GRIS)

out = "EXPLICACION - Memoria de Calculo - Cantidades de Obra.docx"
doc.save(out)
print("DOCUMENTO GUARDADO:", out, "|", os.path.getsize(out)//1024, "KB")
